"""
Unified Analytics Dashboard
----------------------------
Combines Aimfox LinkedIn analytics + Smartlead email analytics
into a single live Flask dashboard.

Run:  python unified_dashboard.py
Open: http://localhost:8080
"""

import json, os, sys, logging, time, threading
from collections import defaultdict
from datetime import datetime

from flask import Flask, jsonify, render_template_string, request, send_file
from dotenv import load_dotenv

load_dotenv()
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from aimfox_client import AimfoxClient
from subsequence_analytics import (
    get_all_campaigns, get_campaign_stats, collect_leads_detail,
    compute_analytics, enrich_replied_leads, resolve_client,
)

try:
    import zoho_sync
    ZOHO_AVAILABLE = True
except Exception:
    ZOHO_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__)
PORT = int(os.getenv("DASHBOARD_PORT", 8081))

LI_CACHE      = {"data": None, "ts": 0}
EMAIL_CACHE   = {"data": None, "ts": 0}
LI_TTL        = 300   # 5 minutes (was 30s)
EMAIL_TTL     = 1800  # 30 minutes (fast parallel fetch makes long cache fine)
REFRESH_EVERY = 1500  # 25 minutes — pre-emptive refresh before TTL expires
_email_lock   = threading.Lock()   # prevent concurrent Smartlead fetches
_email_fetching = False            # True while a fetch is in progress
_email_progress = {"phase": "Idle", "done": 0, "total": 0}

# ── SQLite persistent cache ──────────────────────────────────────────
import sqlite3
DB_PATH = os.getenv("CACHE_DB_PATH",
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), "cache.db"))


def _db_init():
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute("""CREATE TABLE IF NOT EXISTS kv (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL,
            ts INTEGER NOT NULL
        )""")
        conn.commit()
    finally:
        conn.close()


def _db_save(key, value):
    """Atomically persist a JSON-serialisable value under `key`."""
    try:
        conn = sqlite3.connect(DB_PATH, timeout=10)
        conn.execute("INSERT OR REPLACE INTO kv (key, value, ts) VALUES (?, ?, ?)",
                     (key, json.dumps(value, default=str), int(time.time())))
        conn.commit()
        conn.close()
    except Exception as e:
        log.warning("[db] save %s failed: %s", key, e)


def _db_load(key):
    """Return (value, ts) or (None, 0)."""
    try:
        conn = sqlite3.connect(DB_PATH, timeout=10)
        cur  = conn.execute("SELECT value, ts FROM kv WHERE key = ?", (key,))
        row  = cur.fetchone()
        conn.close()
        if row:
            return json.loads(row[0]), int(row[1])
    except Exception as e:
        log.warning("[db] load %s failed: %s", key, e)
    return None, 0


_db_init()
# Hydrate caches from SQLite on startup so dashboard is instant after restart
_e_data, _e_ts = _db_load("email_cache")
if _e_data:
    EMAIL_CACHE["data"] = _e_data
    EMAIL_CACHE["ts"]   = _e_ts
    log.info("[db] Loaded email cache from SQLite (age=%ds, %d parents, %d subs)",
             int(time.time()) - _e_ts,
             len(_e_data.get("parent_analytics", [])),
             len(_e_data.get("sub_analytics", [])))

SAMPLE_EMAIL_DATA = {
    "loading": True,
    "parent_analytics": [
        {"client": "Sample Client A", "status": "ACTIVE", "total": 120, "opened": 72,
         "clicked": 18, "replied": 9, "bounced": 3, "unsubscribed": 2,
         "open_rate": 60.0, "click_rate": 15.0, "reply_rate": 7.5, "positive_rate": 4.0,
         "added_to_sub": 30, "added_to_sub_rate": 25.0, "leads": []},
        {"client": "Sample Client B", "status": "ACTIVE", "total": 95, "opened": 52,
         "clicked": 11, "replied": 6, "bounced": 2, "unsubscribed": 1,
         "open_rate": 54.7, "click_rate": 11.6, "reply_rate": 6.3, "positive_rate": 3.2,
         "added_to_sub": 20, "added_to_sub_rate": 21.0, "leads": []},
    ],
    "sub_analytics": [
        {"parent": "Sample Client A", "subsequence": "Sub-Sequence 1", "status": "ACTIVE",
         "total": 30, "opened": 18, "clicked": 6, "replied": 3,
         "open_rate": 60.0, "click_rate": 20.0, "reply_rate": 10.0, "leads": []},
        {"parent": "Sample Client B", "subsequence": "Sub-Sequence 1", "status": "ACTIVE",
         "total": 20, "opened": 11, "clicked": 3, "replied": 2,
         "open_rate": 55.0, "click_rate": 15.0, "reply_rate": 10.0, "leads": []},
    ],
}


# ── LinkedIn helpers ──────────────────────────────────────────────────

def _owners_list(v):
    if not v: return []
    if isinstance(v, list): return [str(x) for x in v]
    return [x.strip() for x in str(v).split(",") if x.strip()]

def _ts(ms):
    try:
        if isinstance(ms, (int, float)):
            return datetime.fromtimestamp(ms / 1000).strftime("%Y-%m-%d")
        return str(ms)[:10]
    except Exception:
        return str(ms)

def _build_campaign_stats(campaigns, recent_leads):
    acc_per, rep_per = defaultdict(int), defaultdict(int)
    for lead in recent_leads:
        cid = lead.get("campaign_id")
        t   = lead.get("transition", "")
        if t == "accepted": acc_per[cid] += 1
        elif t == "reply":  rep_per[cid] += 1
    rows = []
    for c in campaigns:
        cid        = c.get("id")
        completion = c.get("completion") or 0
        rows.append({
            "id":              cid,
            "name":            c.get("name", "-"),
            "state":           c.get("state", "-"),
            "type":            c.get("type", "-"),
            "created":         _ts(c.get("created_at")),
            "targets":         c.get("target_count") or 0,
            "completion_pct":  f"{completion*100:.0f}%" if isinstance(completion, float) else f"{completion}%",
            "accepted_recent": acc_per.get(cid, 0),
            "replies_recent":  rep_per.get(cid, 0),
            "owners":          _owners_list(c.get("owners")),
        })
    return rows

def _build_li_stats(accounts, campaigns, recent_leads, convos, filter_id=None):
    acc_map        = {a["id"]: a for a in accounts}
    f_camps        = [c for c in campaigns if not filter_id or filter_id in _owners_list(c.get("owners"))]
    f_convos       = [c for c in convos    if not filter_id or str(c.get("owner","")) == filter_id]
    connected      = [c for c in f_convos  if c.get("connected")]
    targets        = sum(c.get("targets",0) or 0 for c in f_camps)
    accepted       = sum(c.get("accepted_recent",0)  for c in f_camps)
    replies        = sum(c.get("replies_recent",0)   for c in f_camps)
    messages_sent  = 0
    for conv in connected:
        acc_name = acc_map.get(conv.get("owner",""), {}).get("full_name","")
        for msg in conv.get("_messages",[]):
            sender = (msg.get("sender") or {}).get("full_name","")
            if msg.get("automated") or sender == acc_name:
                messages_sent += 1
    return {
        "targets_sent":         targets,
        "accepted":             accepted,
        "replies":              replies,
        "messages_sent":        messages_sent,
        "active_campaigns":     sum(1 for c in f_camps  if c.get("state") == "ACTIVE"),
        "total_campaigns":      len(f_camps),
        "active_conversations": len(connected),
        "unread_messages":      sum(c.get("unread_count",0) for c in f_convos),
        "accept_rate":          f"{accepted/targets*100:.1f}%" if targets  else "—",
        "reply_rate":           f"{replies/accepted*100:.1f}%" if accepted else "—",
    }

def get_li_data(force=False):
    now = time.time()
    if not force and LI_CACHE["data"] and (now - LI_CACHE["ts"]) < LI_TTL:
        return LI_CACHE["data"]
    client    = AimfoxClient()
    accounts  = client.list_accounts()
    campaigns = client.list_campaigns()
    recent    = client.get_recent_leads()
    convos    = []
    for conv in client.list_conversations():
        owner, urn = conv.get("owner"), conv.get("conversation_urn")
        try:    conv["_messages"] = client.get_conversation_messages(owner, urn) if owner and urn else []
        except: conv["_messages"] = []
        convos.append(conv)
    rows = _build_campaign_stats(campaigns, recent)
    LI_CACHE.update({"data": (accounts, rows, recent, convos), "ts": now})
    return LI_CACHE["data"]

FETCH_WORKERS = 4    # parallel threads for campaign stats fetching (lowered to avoid Smartlead rate-limits)

def _fetch_sub(sub, parent_map):
    """Fetch one subsequence — runs in a thread pool."""
    parent      = parent_map.get(sub["parent_campaign_id"], {})
    parent_name = resolve_client(parent.get("name", f"ID:{sub['parent_campaign_id']}"))
    stats       = get_campaign_stats(sub["id"])
    leads       = collect_leads_detail(stats)
    a           = compute_analytics(leads)
    emails      = {l["email"] for l in leads}
    return {
        "pid":    sub["parent_campaign_id"],
        "row":    {"parent": parent_name, "subsequence": sub["name"],
                   "status": sub.get("status", "UNKNOWN"), **a, "leads": leads},
        "emails": emails,
    }

def _fetch_parent(pid, parent, sub_emails):
    """Fetch one parent campaign — runs in a thread pool (no message enrichment for speed)."""
    client_name = resolve_client(parent.get("name", f"ID:{pid}"))
    stats       = get_campaign_stats(pid)
    leads       = collect_leads_detail(stats)
    # Skip enrich_replied_leads — message history fetch is very slow (30+ sec per campaign)
    # Message threads are still available via double-click modal when cache warms up
    a           = compute_analytics(leads)
    added       = len(sub_emails.get(pid, set()))
    return {
        "client": client_name, "raw_name": parent.get("name", ""),
        "status": parent.get("status", "UNKNOWN"), **a,
        "added_to_sub":      added,
        "added_to_sub_rate": round(added / a["total"] * 100, 2) if a["total"] else 0,
        "leads": leads,
    }

def _do_email_fetch():
    """Run the full Smartlead fetch in parallel and populate EMAIL_CACHE."""
    global _email_fetching, _email_progress
    from concurrent.futures import ThreadPoolExecutor, as_completed
    log.info("Fetching Smartlead email data (parallel, %d workers)…", FETCH_WORKERS)
    try:
        now = time.time()
        _email_progress = {"phase": "Loading campaigns…", "done": 0, "total": 0}
        all_cmp    = get_all_campaigns()
        parent_map = {c["id"]: c for c in all_cmp}
        p_to_subs  = defaultdict(list)
        for c in all_cmp:
            if c.get("parent_campaign_id"):
                p_to_subs[c["parent_campaign_id"]].append(c)
        parents_with_subs = {pid: parent_map[pid] for pid in p_to_subs if pid in parent_map}

        # ── PHASE 1: subsequences (parallel) ────────────────────────
        subs = [c for c in all_cmp if c.get("parent_campaign_id")]
        _email_progress = {"phase": "Fetching subsequence stats…", "done": 0, "total": len(subs)}
        sub_emails  = defaultdict(set)
        sub_data    = []
        done_count  = 0
        with ThreadPoolExecutor(max_workers=FETCH_WORKERS) as ex:
            futs = {ex.submit(_fetch_sub, sub, parent_map): sub for sub in subs}
            for fut in as_completed(futs):
                done_count += 1
                _email_progress["done"] = done_count
                try:
                    r = fut.result()
                    sub_data.append(r["row"])
                    sub_emails[r["pid"]].update(r["emails"])
                except Exception as e:
                    log.warning("Sub fetch error: %s", e)

        # ── PHASE 2: parent campaigns (parallel) ─────────────────────
        parents    = list(parents_with_subs.items())
        _email_progress = {"phase": "Fetching parent campaign stats…", "done": 0, "total": len(parents)}
        parent_data = []
        done_count  = 0
        with ThreadPoolExecutor(max_workers=FETCH_WORKERS) as ex:
            futs = {ex.submit(_fetch_parent, pid, parent, sub_emails): pid
                    for pid, parent in parents}
            for fut in as_completed(futs):
                done_count += 1
                _email_progress["done"] = done_count
                try:
                    parent_data.append(fut.result())
                except Exception as e:
                    log.warning("Parent fetch error: %s", e)

        result = {"parent_analytics": parent_data, "sub_analytics": sub_data}
        EMAIL_CACHE.update({"data": result, "ts": now})
        _email_progress = {"phase": "Done", "done": len(parents), "total": len(parents)}
        log.info("Email data fetched in %.1fs (%d parents, %d subs).",
                 time.time() - now, len(parent_data), len(sub_data))
        # Persist to SQLite so we don't lose the cache on restart
        _db_save("email_cache", result)
    finally:
        _email_fetching = False


def get_email_data(force=False):
    global _email_fetching
    now = time.time()
    # Return cached data if still fresh
    if not force and EMAIL_CACHE["data"] and (now - EMAIL_CACHE["ts"]) < EMAIL_TTL:
        return EMAIL_CACHE["data"]
    # If another thread is already fetching, wait for it to finish then return cache
    if _email_fetching and not force:
        log.info("Email fetch already in progress — waiting for it to complete...")
        with _email_lock:   # blocks until the running fetch releases it
            pass
        return EMAIL_CACHE["data"]
    # Acquire lock, set flag, run fetch (flag cleared in finally inside _do_email_fetch)
    with _email_lock:
        # Re-check cache after acquiring lock (another thread may have just filled it)
        now = time.time()
        if not force and EMAIL_CACHE["data"] and (now - EMAIL_CACHE["ts"]) < EMAIL_TTL:
            return EMAIL_CACHE["data"]
        _email_fetching = True
        _do_email_fetch()
    return EMAIL_CACHE["data"]


# ── API endpoints ─────────────────────────────────────────────────────

@app.route("/api/linkedin")
def api_linkedin():
    try:
        accounts, campaigns, recent, convos = get_li_data()
        acc_map = {a["id"]: a.get("full_name", a["id"]) for a in accounts}
        account_stats = []
        for acc in accounts:
            s = _build_li_stats(accounts, campaigns, recent, convos, acc["id"])
            account_stats.append({"id": acc["id"], "name": acc.get("full_name", acc["id"]),
                                   "picture": acc.get("picture_url",""), "occupation": acc.get("occupation",""),
                                   "premium": acc.get("premium", False), "state": acc.get("state",""), "stats": s})
        camp_list = [{**{k:v for k,v in c.items() if k!="owners"},
                      "owners": c["owners"], "owner_names": [acc_map.get(o,o) for o in c["owners"]]}
                     for c in campaigns]
        leads_list = [{"transition": e.get("transition",""), "timestamp": (e.get("timestamp") or "")[:10],
                       "campaign_name": e.get("campaign_name",""), "account_id": e.get("account_id",""),
                       "account_name": acc_map.get(e.get("account_id",""),""),
                       "target_name": (e.get("target") or {}).get("full_name",""),
                       "target_occupation": (e.get("target") or {}).get("occupation","")} for e in recent]
        conv_list = []
        for conv in convos:
            parts = conv.get("participants",[])
            lead  = parts[0] if parts else {}
            msgs  = [{"sender": (m.get("sender") or {}).get("full_name",""),
                      "body": ((m.get("body") or "").strip() or "[attachment]")[:300],
                      "automated": m.get("automated", False),
                      "date": str(m.get("created_at",""))[:10]}
                     for m in conv.get("_messages",[]) if (m.get("body") or "").strip() or m.get("attachments")]
            conv_list.append({"owner": conv.get("owner",""), "owner_name": acc_map.get(conv.get("owner",""),""),
                               "connected": conv.get("connected", False), "unread": conv.get("unread_count",0),
                               "lead_name": lead.get("full_name",""), "lead_picture": lead.get("picture_url",""),
                               "lead_occupation": lead.get("occupation",""), "messages": msgs})
        return jsonify({"overall": _build_li_stats(accounts, campaigns, recent, convos),
                        "accounts": account_stats, "campaigns": camp_list,
                        "recent_leads": leads_list, "conversations": conv_list,
                        "cached": (time.time() - LI_CACHE["ts"]) < LI_TTL})
    except Exception as e:
        log.error("LinkedIn API error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route("/api/email")
def api_email():
    try:
        # Return sample data immediately if fetch is in progress and cache is empty
        if _email_fetching and not EMAIL_CACHE["data"]:
            return jsonify({**SAMPLE_EMAIL_DATA,
                            "loading": True,
                            "progress": _email_progress,
                            "cached": False})
        data = get_email_data()
        if data is None:
            return jsonify({**SAMPLE_EMAIL_DATA, "loading": True,
                            "progress": _email_progress, "cached": False})
        return jsonify({**data, "loading": False,
                        "cached": (time.time() - EMAIL_CACHE["ts"]) < EMAIL_TTL})
    except Exception as e:
        log.error("Email API error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route("/api/email/status")
def api_email_status():
    return jsonify({
        "loading": _email_fetching,
        "progress": _email_progress,
        "cached": bool(EMAIL_CACHE["data"]),
        "cache_age": round(time.time() - EMAIL_CACHE["ts"]) if EMAIL_CACHE["ts"] else None,
    })


@app.route("/api/cache/info")
def api_cache_info():
    """Diagnostic endpoint — see freshness of all caches at once."""
    now = time.time()
    return jsonify({
        "email": {
            "cached":      bool(EMAIL_CACHE["data"]),
            "age_seconds": round(now - EMAIL_CACHE["ts"]) if EMAIL_CACHE["ts"] else None,
            "ttl_seconds": EMAIL_TTL,
            "fetching":    _email_fetching,
            "parents":     len((EMAIL_CACHE["data"] or {}).get("parent_analytics", [])),
            "subs":        len((EMAIL_CACHE["data"] or {}).get("sub_analytics", [])),
        },
        "linkedin": {
            "cached":      bool(LI_CACHE["data"]),
            "age_seconds": round(now - LI_CACHE["ts"]) if LI_CACHE["ts"] else None,
            "ttl_seconds": LI_TTL,
        },
        "scheduler": {
            "refresh_every_seconds": REFRESH_EVERY,
            "next_refresh_estimate": (round((EMAIL_CACHE["ts"] + REFRESH_EVERY) - now)
                                       if EMAIL_CACHE["ts"] else None),
        },
        "db_path": DB_PATH,
    })


@app.route("/api/cache/refresh", methods=["POST"])
def api_cache_refresh():
    """Manual refresh trigger — kicks off a background fetch and returns immediately."""
    if _email_fetching:
        return jsonify({"ok": False, "error": "Already fetching"}), 409
    threading.Thread(target=lambda: get_email_data(force=True), daemon=True).start()
    return jsonify({"ok": True, "message": "Refresh started in background"})

@app.route("/api/zoho-stats")
def api_zoho_stats():
    path = os.path.join(os.path.dirname(__file__), "reports", "zoho_sync_stats.json")
    try:
        with open(path) as f: return jsonify(json.load(f))
    except FileNotFoundError:
        return jsonify({"last_sync": None, "created": 0, "updated": 0, "total": 0, "by_company": {}})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/zoho-sync", methods=["POST"])
def api_zoho_sync():
    if not ZOHO_AVAILABLE:
        return jsonify({"ok": False, "error": "Zoho credentials not configured"}), 400
    try:
        counts = zoho_sync.run_sync(sync_conversations=False, module="Leads")
        return jsonify({"ok": True, **counts})
    except Exception as e:
        log.error("Zoho sync error: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500

# ── Helpers for export / email ────────────────────────────────────────

def _safe_text(s):
    """Sanitise string for fpdf2 latin-1 core fonts.
    Substitute common Unicode chars that the core font can't render
    (em-dash, middle dot, arrows, smart quotes) with ASCII equivalents
    before falling back to latin-1 with replace."""
    if not s:
        return ""
    s = str(s)
    repl = {
        "—": "-",   # em-dash —
        "–": "-",   # en-dash –
        "·": "|",   # middle dot ·
        "•": "*",   # bullet •
        "→": "->",  # right arrow →
        "←": "<-",  # left arrow ←
        "…": "...", # ellipsis …
        "‘": "'",   # left single quote ‘
        "’": "'",   # right single quote ’
        "“": '"',   # left double quote “
        "”": '"',   # right double quote ”
        " ": " ",   # nbsp
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")

def _get_lead_last_activity(lead):
    dates = [d for d in [lead.get("reply_time"), lead.get("open_time"),
                          lead.get("click_time"), lead.get("sent_time")] if d]
    return max(dates)[:10] if dates else ""

def _get_lead_summary(lead):
    msgs = lead.get("messages", [])
    if not msgs:
        return ""
    return ((msgs[-1].get("body") or "")[:150])

def _pct(num, den):
    return f"{(num/den*100):.1f}%" if den else "—"


def _compute_date_range(parent_data, sub_data):
    """Find earliest sent_time and latest activity time across all leads."""
    times = []
    for collection in (parent_data or [], sub_data or []):
        for row in collection:
            for l in row.get("leads", []):
                for k in ("sent_time", "open_time", "click_time", "reply_time"):
                    v = l.get(k)
                    if v and isinstance(v, str) and len(v) >= 10:
                        times.append(v[:10])
    if not times:
        return None, None
    return min(times), max(times)


def _build_summary_paragraph(client_name, summary, parent_data, sub_data, pos_leads):
    """Generate a 2–3 sentence executive summary using the actual numbers."""
    label = client_name or "All Clients"
    date_from, date_to = _compute_date_range(parent_data, sub_data)
    range_str = (f" (data from {date_from} to {date_to})"
                 if date_from and date_to and date_from != date_to
                 else (f" (data on {date_from})" if date_from else ""))
    sent          = summary.get("sent", 0)
    opened        = summary.get("opened", 0)
    replied       = summary.get("replied", 0)
    added_sub     = summary.get("added_to_sub", 0)
    sub_total     = summary.get("sub_total", 0)
    sub_opened    = summary.get("sub_opened", 0)
    sub_replied   = summary.get("sub_replied", 0)
    positive      = summary.get("positive", len(pos_leads or []))
    bounced       = summary.get("bounced", 0)
    n_camps       = len(parent_data or [])
    n_subs        = len(sub_data or [])

    s1 = (f"{label} ran {n_camps} parent campaign{'s' if n_camps != 1 else ''}{range_str} "
          f"with {sent:,} emails sent — {opened:,} opens ({_pct(opened, sent)}) "
          f"and {replied:,} replies ({_pct(replied, sent)}).")
    s2 = (f"{added_sub:,} leads ({_pct(added_sub, sent)}) were added to {n_subs} "
          f"subsequence{'s' if n_subs != 1 else ''}, where {sub_opened:,} "
          f"opened ({_pct(sub_opened, sub_total)}) and {sub_replied:,} replied "
          f"({_pct(sub_replied, sub_total)}).")
    s3 = (f"{positive:,} positive repl{'ies' if positive != 1 else 'y'} "
          f"identified · {bounced:,} bounced.")
    return " ".join([s1, s2, s3])


def _compute_summary(parent_data, sub_data):
    def s(rows, key): return sum((r.get(key) or 0) for r in rows)
    sent = s(parent_data, "total")
    sub_total = s(sub_data, "total")
    return {
        "sent":         sent,
        "opened":       s(parent_data, "opened"),
        "clicked":      s(parent_data, "clicked"),
        "replied":      s(parent_data, "replied"),
        "added_to_sub": s(parent_data, "added_to_sub"),
        "bounced":      s(parent_data, "bounced"),
        "unsubscribed": s(parent_data, "unsubscribed"),
        "positive":     s(parent_data, "positive"),
        "sub_total":    sub_total,
        "sub_opened":   s(sub_data, "opened"),
        "sub_clicked":  s(sub_data, "clicked"),
        "sub_replied":  s(sub_data, "replied"),
    }


def _build_report_pdf(parent_data, sub_data, client_name=None,
                      summary=None, positive_leads=None):
    """PDF report matching the on-screen layout: header → summary cards
    → parent campaigns → subsequences → positive leads."""
    from fpdf import FPDF
    import io

    if summary is None:
        summary = _compute_summary(parent_data, sub_data)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # ── Header ──────────────────────────────────────────────────────
    title = (f"{client_name} — Client Dashboard Report"
             if client_name else "Smartlead Analytics Report")
    pdf.set_fill_color(31, 27, 75)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, _safe_text(title), ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)
    date_from, date_to = _compute_date_range(parent_data, sub_data)
    range_meta = (f"   ·   Data range: {date_from} → {date_to}"
                  if date_from and date_to else "")
    pdf.cell(0, 6, _safe_text(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                              + (f"   ·   Campaigns: {len(parent_data)}   ·   Subsequences: {len(sub_data)}" if client_name else "")
                              + range_meta),
             ln=True)
    pdf.ln(2)

    # ── 2-3 line executive summary ──────────────────────────────────
    paragraph = _build_summary_paragraph(client_name, summary, parent_data, sub_data, positive_leads)
    pdf.set_fill_color(238, 242, 255)
    pdf.set_draw_color(199, 210, 254)
    y = pdf.get_y()
    # Reserve box height; multi_cell auto-wraps
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(67, 56, 202)
    pdf.cell(0, 6, _safe_text("Executive Summary"), ln=True, fill=True, border="LTR")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(30, 41, 59)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 5, _safe_text(paragraph), border="LBR", fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    # ── Summary cards (11 stats in a 4-col grid) ────────────────────
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, _safe_text("Summary Overview"), ln=True)
    pdf.ln(1)

    cards = [
        ("Sent",          summary.get("sent", 0),         None),
        ("Opened",        summary.get("opened", 0),       _pct(summary.get("opened", 0),  summary.get("sent", 0))),
        ("Clicked",       summary.get("clicked", 0),      _pct(summary.get("clicked", 0), summary.get("sent", 0))),
        ("Replied",       summary.get("replied", 0),      _pct(summary.get("replied", 0), summary.get("sent", 0))),
        ("Added to Sub",  summary.get("added_to_sub", 0), _pct(summary.get("added_to_sub", 0), summary.get("sent", 0))),
        ("Sub Opened",    summary.get("sub_opened", 0),   _pct(summary.get("sub_opened", 0),   summary.get("sub_total", 0))),
        ("Sub Clicked",   summary.get("sub_clicked", 0),  _pct(summary.get("sub_clicked", 0),  summary.get("sub_total", 0))),
        ("Sub Replied",   summary.get("sub_replied", 0),  _pct(summary.get("sub_replied", 0),  summary.get("sub_total", 0))),
        ("Positive Replies", summary.get("positive", 0),  None),
        ("Bounced",       summary.get("bounced", 0),      None),
        ("Unsubscribed",  summary.get("unsubscribed", 0), None),
    ]
    card_w, card_h = 45, 18
    cols = 4
    for i, (lbl, val, sub) in enumerate(cards):
        if i and i % cols == 0:
            pdf.ln(card_h + 1)
        x = pdf.get_x()
        y = pdf.get_y()
        # Card background
        pdf.set_fill_color(248, 250, 252)
        pdf.set_draw_color(226, 232, 240)
        pdf.rect(x, y, card_w, card_h, style="DF")
        # Big value
        pdf.set_xy(x + 2, y + 2)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(card_w - 4, 6, _safe_text(f"{val:,}" if isinstance(val, (int, float)) else str(val)))
        # Label
        pdf.set_xy(x + 2, y + 9)
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(100, 116, 139)
        pdf.cell(card_w - 4, 4, _safe_text(lbl))
        # Sub (rate)
        if sub:
            pdf.set_xy(x + 2, y + 13)
            pdf.set_font("Helvetica", "B", 7)
            pdf.set_text_color(79, 70, 229)
            pdf.cell(card_w - 4, 4, _safe_text(sub))
        pdf.set_xy(x + card_w + 1, y)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(card_h + 6)

    # ── Parent Campaigns table ──────────────────────────────────────
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, _safe_text("Parent Campaigns"), ln=True)
    pdf.ln(1)
    p_hdrs = ["Campaign", "Status", "Total", "Opened", "Clicked", "Replied",
              "Bounced", "Added Sub", "Open %", "Reply %"]
    p_w    = [60, 18, 14, 16, 16, 16, 16, 18, 14, 14]
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(241, 245, 249)
    for h, w in zip(p_hdrs, p_w):
        pdf.cell(w, 7, _safe_text(h), border=1, fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 7)
    for r in parent_data:
        vals = [
            (r.get("raw_name") or r.get("client", ""))[:42],
            r.get("status", "—"),
            f"{r.get('total', 0):,}",
            f"{r.get('opened', 0):,}",
            f"{r.get('clicked', 0):,}",
            f"{r.get('replied', 0):,}",
            f"{r.get('bounced', 0):,}",
            f"{r.get('added_to_sub', 0):,}",
            f"{(r.get('open_rate') or 0):.1f}%",
            f"{(r.get('reply_rate') or 0):.1f}%",
        ]
        for v, w in zip(vals, p_w):
            pdf.cell(w, 5.5, _safe_text(str(v)), border=1)
        pdf.ln()

    # ── Subsequences table (only those with leads) ──────────────────
    active_subs = [s for s in (sub_data or []) if (s.get("total") or 0) > 0]
    if active_subs:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, _safe_text(f"Subsequences ({len(active_subs)} active of {len(sub_data)})"), ln=True)
        pdf.ln(1)
        s_hdrs = ["Subsequence", "Total", "Opened", "Clicked", "Replied", "Open %", "Reply %"]
        s_w    = [80, 18, 18, 18, 18, 18, 18]
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(241, 245, 249)
        for h, w in zip(s_hdrs, s_w):
            pdf.cell(w, 7, _safe_text(h), border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 7)
        # Sort by total desc so biggest subs come first
        active_subs.sort(key=lambda r: -(r.get("total") or 0))
        for r in active_subs:
            vals = [
                (r.get("subsequence") or r.get("parent", ""))[:48],
                f"{r.get('total', 0):,}",
                f"{r.get('opened', 0):,}",
                f"{r.get('clicked', 0):,}",
                f"{r.get('replied', 0):,}",
                f"{(r.get('open_rate') or 0):.1f}%",
                f"{(r.get('reply_rate') or 0):.1f}%",
            ]
            for v, w in zip(vals, s_w):
                pdf.cell(w, 5.5, _safe_text(str(v)), border=1)
            pdf.ln()

    # ── Positive Leads ──────────────────────────────────────────────
    if positive_leads:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, _safe_text(f"Positive Leads ({len(positive_leads)})"), ln=True)
        pdf.ln(1)
        l_hdrs = ["Lead", "Email", "Category", "Reply Date", "Campaign"]
        l_w    = [42, 60, 24, 22, 54]
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(241, 245, 249)
        for h, w in zip(l_hdrs, l_w):
            pdf.cell(w, 7, _safe_text(h), border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 7)
        for l in positive_leads[:60]:  # cap to keep PDF reasonable
            vals = [
                (l.get("name") or "")[:30],
                (l.get("email") or "")[:42],
                (l.get("category") or l.get("status") or "")[:18],
                (l.get("reply_time") or l.get("last_activity") or "")[:10],
                (l.get("campaign") or l.get("source") or "")[:38],
            ]
            for v, w in zip(vals, l_w):
                pdf.cell(w, 5.5, _safe_text(str(v)), border=1)
            pdf.ln()
        if len(positive_leads) > 60:
            pdf.set_font("Helvetica", "I", 7)
            pdf.cell(0, 5, _safe_text(f"… {len(positive_leads) - 60} more leads not shown"), ln=True)

    return io.BytesIO(pdf.output())


# ── New API routes ─────────────────────────────────────────────────────

@app.route("/api/linkedin-stats")
def api_linkedin_stats():
    """Filtered LinkedIn KPI stats for the segment stat card."""
    try:
        segment     = request.args.get("segment", "all")
        sub_segment = request.args.get("sub_segment", "all")
        accounts, campaigns, recent, convos = get_li_data()

        client  = request.args.get("client", "ALL")

        # Map email client name → LinkedIn campaign keyword (same as CLIENT_MAP)
        CLIENT_KEYWORDS = {
            "FEAAM": "FEAAM", "Wastestream": "WASTESTREAM", "KM": "KM",
            "Stanford G": "STANFORD", "Nexus": "NEXUS",
            "Henig": "HENIG", "Kendra": "USA",
        }

        f_camps = campaigns[:]
        if client != "ALL" and client in CLIENT_KEYWORDS:
            kw = CLIENT_KEYWORDS[client]
            f_camps = [c for c in f_camps if kw in (c.get("name") or "").upper()]
        if segment != "all":
            f_camps = [c for c in f_camps if segment in _owners_list(c.get("owners"))]
        if sub_segment != "all":
            f_camps = [c for c in f_camps
                       if (c.get("type") or "").upper() == sub_segment.upper()]

        camp_ids  = {c["id"] for c in f_camps}
        f_leads   = [l for l in recent if l.get("campaign_id") in camp_ids]
        targets   = sum(c.get("targets", 0) or 0 for c in f_camps)
        accepted  = sum(1 for l in f_leads if l.get("transition") == "accepted")
        replies   = sum(1 for l in f_leads if l.get("transition") == "reply")

        seg_opts = [{"id": a["id"], "name": a.get("full_name", a["id"])} for a in accounts]
        types    = sorted({c.get("type") or "" for c in campaigns if c.get("type")})
        sub_opts = [{"id": t, "name": t} for t in types]

        return jsonify({
            "total_leads":      len(f_leads),
            "requests_sent":    targets,
            "accepted":         accepted,
            "replies":          replies,
            "acceptance_rate":  f"{accepted/targets*100:.1f}%" if targets  else "—",
            "reply_rate":       f"{replies/accepted*100:.1f}%" if accepted else "—",
            "segment_opts":     seg_opts,
            "sub_segment_opts": sub_opts,
        })
    except Exception as e:
        log.error("LinkedIn stats error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/positive-leads")
def api_positive_leads():
    """Returns positive leads from the cached email data."""
    try:
        data = EMAIL_CACHE.get("data")
        if not data:
            return jsonify({"leads": [], "loading": True})
        POSITIVE = {"interested", "meeting booked", "positive",
                    "meeting request", "will buy", "warm", "demo request"}
        leads = []
        for campaign in data.get("parent_analytics", []):
            for l in campaign.get("leads", []):
                if (l.get("category") or "").lower() in POSITIVE:
                    leads.append({
                        "name":          l.get("name") or l.get("email") or "Unknown",
                        "company":       campaign.get("client", ""),
                        "source":        campaign.get("raw_name", ""),
                        "summary":       _get_lead_summary(l),
                        "status":        l.get("category", ""),
                        "last_activity": _get_lead_last_activity(l),
                    })
        return jsonify({"leads": leads, "loading": False})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _extract_export_args(body):
    """Pull the per-client filtered payload out of the request body, with cache fallback."""
    cache       = EMAIL_CACHE.get("data") or {}
    parent_data = body.get("parent_analytics", cache.get("parent_analytics", []))
    sub_data    = body.get("sub_analytics",    cache.get("sub_analytics",    []))
    client_name = body.get("client") or None
    summary     = body.get("summary") or _compute_summary(parent_data, sub_data)
    pos_leads   = body.get("positive_leads") or []
    return parent_data, sub_data, client_name, summary, pos_leads


@app.route("/api/export/pdf", methods=["POST"])
def api_export_pdf():
    try:
        body = request.get_json(silent=True) or {}
        parent_data, sub_data, client_name, summary, pos_leads = _extract_export_args(body)
        buf = _build_report_pdf(parent_data, sub_data,
                                client_name=client_name,
                                summary=summary,
                                positive_leads=pos_leads)
        slug = (client_name or "all").lower().replace(" ", "_")
        filename = f"{slug}_analytics_{datetime.now().strftime('%Y%m%d')}.pdf"
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True, download_name=filename)
    except ImportError:
        return jsonify({"error": "fpdf2 not installed — run: pip install fpdf2"}), 500
    except Exception as e:
        log.error("PDF export error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/export/docx", methods=["POST"])
def api_export_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        body = request.get_json(silent=True) or {}
        parent_data, sub_data, client_name, summary, pos_leads = _extract_export_args(body)

        doc = Document()
        title = (f"{client_name} — Client Dashboard Report"
                 if client_name else "Smartlead Analytics Report")
        h = doc.add_heading(title, 0)
        sub = doc.add_paragraph()
        sub.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
        if client_name:
            sub.add_run(f"   ·   Campaigns: {len(parent_data)}   ·   Subsequences: {len(sub_data)}").italic = True
        df, dt = _compute_date_range(parent_data, sub_data)
        if df and dt:
            sub.add_run(f"   ·   Data range: {df} → {dt}").italic = True

        # ── Executive Summary (2-3 sentences) ────────────────────────
        doc.add_heading("Executive Summary", level=1)
        para = doc.add_paragraph(
            _build_summary_paragraph(client_name, summary, parent_data, sub_data, pos_leads)
        )
        para.runs[0].font.size = Pt(11)

        # ── Summary ─────────────────────────────────────────────────
        doc.add_heading("Summary Overview", level=1)
        sum_tbl = doc.add_table(rows=1, cols=3)
        sum_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Metric", "Value", "Rate"]):
            sum_tbl.rows[0].cells[i].text = h
        rows = [
            ("Sent",          summary.get("sent", 0),         "—"),
            ("Opened",        summary.get("opened", 0),       _pct(summary.get("opened", 0),       summary.get("sent", 0))),
            ("Clicked",       summary.get("clicked", 0),      _pct(summary.get("clicked", 0),      summary.get("sent", 0))),
            ("Replied",       summary.get("replied", 0),      _pct(summary.get("replied", 0),      summary.get("sent", 0))),
            ("Added to Sub",  summary.get("added_to_sub", 0), _pct(summary.get("added_to_sub", 0), summary.get("sent", 0))),
            ("Sub Opened",    summary.get("sub_opened", 0),   _pct(summary.get("sub_opened", 0),   summary.get("sub_total", 0))),
            ("Sub Clicked",   summary.get("sub_clicked", 0),  _pct(summary.get("sub_clicked", 0),  summary.get("sub_total", 0))),
            ("Sub Replied",   summary.get("sub_replied", 0),  _pct(summary.get("sub_replied", 0),  summary.get("sub_total", 0))),
            ("Positive Replies", summary.get("positive", 0),  "—"),
            ("Bounced",       summary.get("bounced", 0),      "—"),
            ("Unsubscribed",  summary.get("unsubscribed", 0), "—"),
        ]
        for lbl, val, rate in rows:
            cells = sum_tbl.add_row().cells
            cells[0].text = lbl
            cells[1].text = f"{val:,}"
            cells[2].text = rate

        # ── Parent Campaigns ────────────────────────────────────────
        doc.add_heading("Parent Campaigns", level=1)
        tbl = doc.add_table(rows=1, cols=10)
        tbl.style = "Light Grid Accent 1"
        for i, hd in enumerate(["Campaign", "Status", "Total", "Opened", "Clicked",
                                "Replied", "Bounced", "Added to Sub", "Open %", "Reply %"]):
            tbl.rows[0].cells[i].text = hd
        for r in parent_data:
            cells = tbl.add_row().cells
            cells[0].text = str(r.get("raw_name") or r.get("client", ""))
            cells[1].text = str(r.get("status", "—"))
            cells[2].text = f"{r.get('total', 0):,}"
            cells[3].text = f"{r.get('opened', 0):,}"
            cells[4].text = f"{r.get('clicked', 0):,}"
            cells[5].text = f"{r.get('replied', 0):,}"
            cells[6].text = f"{r.get('bounced', 0):,}"
            cells[7].text = f"{r.get('added_to_sub', 0):,}"
            cells[8].text = f"{(r.get('open_rate') or 0):.1f}%"
            cells[9].text = f"{(r.get('reply_rate') or 0):.1f}%"

        # ── Subsequences (only those with leads, sorted by size) ─────
        active_subs = [s for s in (sub_data or []) if (s.get("total") or 0) > 0]
        active_subs.sort(key=lambda r: -(r.get("total") or 0))
        if active_subs:
            doc.add_heading(f"Subsequences ({len(active_subs)} active of {len(sub_data)})", level=1)
            tbl2 = doc.add_table(rows=1, cols=7)
            tbl2.style = "Light Grid Accent 1"
            for i, hd in enumerate(["Subsequence", "Total", "Opened", "Clicked",
                                    "Replied", "Open %", "Reply %"]):
                tbl2.rows[0].cells[i].text = hd
            for r in active_subs:
                cells = tbl2.add_row().cells
                cells[0].text = str(r.get("subsequence") or r.get("parent", ""))
                cells[1].text = f"{r.get('total', 0):,}"
                cells[2].text = f"{r.get('opened', 0):,}"
                cells[3].text = f"{r.get('clicked', 0):,}"
                cells[4].text = f"{r.get('replied', 0):,}"
                cells[5].text = f"{(r.get('open_rate') or 0):.1f}%"
                cells[6].text = f"{(r.get('reply_rate') or 0):.1f}%"

        # ── Positive Leads ──────────────────────────────────────────
        if pos_leads:
            doc.add_heading(f"Positive Leads ({len(pos_leads)})", level=1)
            tbl3 = doc.add_table(rows=1, cols=5)
            tbl3.style = "Light Grid Accent 1"
            for i, hd in enumerate(["Lead", "Email", "Category", "Reply Date", "Campaign"]):
                tbl3.rows[0].cells[i].text = hd
            for l in pos_leads[:200]:
                cells = tbl3.add_row().cells
                cells[0].text = str(l.get("name", ""))
                cells[1].text = str(l.get("email", ""))
                cells[2].text = str(l.get("category") or l.get("status", ""))
                cells[3].text = str((l.get("reply_time") or l.get("last_activity") or ""))[:10]
                cells[4].text = str(l.get("campaign") or l.get("source", ""))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        slug = (client_name or "all").lower().replace(" ", "_")
        filename = f"{slug}_analytics_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=filename,
        )
    except ImportError:
        return jsonify({"error": "python-docx not installed — run: pip install python-docx"}), 500
    except Exception as e:
        log.error("DOCX export error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/share/email", methods=["POST"])
def api_share_email():
    try:
        import smtplib, io
        from email.mime.multipart    import MIMEMultipart
        from email.mime.text         import MIMEText
        from email.mime.application  import MIMEApplication

        body     = request.get_json(silent=True) or {}
        to_email = (body.get("email") or "").strip()
        if not to_email:
            return jsonify({"ok": False, "error": "Email address required"}), 400

        smtp_host = os.getenv("SMTP_HOST", "")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        smtp_user = os.getenv("SMTP_USER", "")
        smtp_pass = os.getenv("SMTP_PASS", "")
        smtp_from = os.getenv("SMTP_FROM", smtp_user)

        if not smtp_host or not smtp_user:
            return jsonify({
                "ok": False,
                "error": "SMTP not configured — add SMTP_HOST, SMTP_USER, SMTP_PASS to .env",
            }), 400

        # Accept per-client filtered data via body, fall back to full cache
        parent_data, sub_data, client_name, summary, pos_leads = _extract_export_args(body)
        client_lbl  = client_name or ""
        pdf_bytes   = _build_report_pdf(parent_data, sub_data,
                                        client_name=client_name,
                                        summary=summary,
                                        positive_leads=pos_leads).read()

        msg            = MIMEMultipart()
        msg["From"]    = smtp_from
        msg["To"]      = to_email
        subj_tag       = f" — {client_lbl}" if client_lbl else ""
        msg["Subject"] = f"Smartlead Analytics Report{subj_tag} — {datetime.now().strftime('%Y-%m-%d')}"
        # Email body opens with the same exec-summary paragraph that's in the PDF
        body_text = (
            (f"Hi,\n\nAttached is the {client_lbl} client dashboard report.\n\n"
             if client_lbl else "Hi,\n\nAttached is the Smartlead analytics report.\n\n")
            + "EXECUTIVE SUMMARY\n"
            + _build_summary_paragraph(client_name, summary, parent_data, sub_data, pos_leads)
            + "\n\nFull breakdown — campaigns, subsequences, positive leads — is in the attached PDF.\n\n"
            + "— Sent from Unified Analytics Dashboard"
        )
        msg.attach(MIMEText(body_text, "plain"))
        fname = f"analytics_{datetime.now().strftime('%Y%m%d')}.pdf"
        part  = MIMEApplication(pdf_bytes, Name=fname)
        part["Content-Disposition"] = f'attachment; filename="{fname}"'
        msg.attach(part)

        with smtplib.SMTP(smtp_host, smtp_port) as srv:
            srv.starttls()
            srv.login(smtp_user, smtp_pass)
            srv.sendmail(smtp_from, to_email, msg.as_string())

        return jsonify({"ok": True})
    except ImportError:
        return jsonify({"ok": False, "error": "fpdf2 not installed — run: pip install fpdf2"}), 500
    except Exception as e:
        log.error("Share email error: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


# ── Per-client sub-dashboards ─────────────────────────────────────────

# Display name → URL slug + LinkedIn campaign keywords (any match counts)
CLIENTS = [
    {"name": "Kendra",      "slug": "kendra",      "li_keywords": ["KENDRA", "USA"]},
    {"name": "FEAAM",       "slug": "feaam",       "li_keywords": ["FEAAM"]},
    {"name": "KM",          "slug": "km",          "li_keywords": ["KM", "K&M"]},
    {"name": "Nexus",       "slug": "nexus",       "li_keywords": ["NEXUS"]},
    {"name": "Henig",       "slug": "henig",       "li_keywords": ["HENIG"]},
    {"name": "Stanford G",  "slug": "stanford",    "li_keywords": ["STANFORD", "STANDFORD"]},
    {"name": "Wastestream", "slug": "wastestream", "li_keywords": ["WASTESTREAM"]},
]
CLIENTS_BY_SLUG = {c["slug"]: c for c in CLIENTS}


def _matches_client(name, keywords):
    up = (name or "").upper()
    return any(kw in up for kw in keywords)


def _slim_lead(l, parent_name=""):
    """Compact lead dict with conversation messages preserved."""
    return {
        "name":     l.get("name") or l.get("email") or "Unknown",
        "email":    l.get("email", ""),
        "category": l.get("category", ""),
        "company":  l.get("company", ""),
        "campaign": parent_name,
        "open_time":  (l.get("open_time")  or "")[:19],
        "click_time": (l.get("click_time") or "")[:19],
        "reply_time": (l.get("reply_time") or "")[:19],
        "sent_time":  (l.get("sent_time")  or "")[:19],
        "summary":  _get_lead_summary(l),
        "messages": [{"from": m.get("from", ""), "type": m.get("type", ""),
                      "time": (m.get("time") or "")[:19],
                      "body": (m.get("body") or "")[:600]}
                     for m in (l.get("messages") or [])][:25],
    }


@app.route("/api/client/<slug>")
def api_client(slug):
    """Bundle of everything a per-client sub-dashboard needs."""
    client = CLIENTS_BY_SLUG.get(slug)
    if not client:
        return jsonify({"error": f"Unknown client '{slug}'"}), 404
    name = client["name"]
    keywords = client["li_keywords"]

    # ── Email (Smartlead) — filter cached parent + sub data by client ─
    email = EMAIL_CACHE.get("data") or {}
    parents = [p for p in email.get("parent_analytics", []) if p.get("client") == name]
    subs    = [s for s in email.get("sub_analytics",    []) if s.get("parent")  == name]

    # Aggregate stat totals (clickable cards)
    def _sum(rows, key): return sum((r.get(key) or 0) for r in rows)
    stats = {
        "sent":          _sum(parents, "total"),
        "opened":        _sum(parents, "opened"),
        "clicked":       _sum(parents, "clicked"),
        "replied":       _sum(parents, "replied"),
        "bounced":       _sum(parents, "bounced"),
        "unsubscribed":  _sum(parents, "unsubscribed"),
        "added_to_sub":  _sum(parents, "added_to_sub"),
        "sub_total":     _sum(subs,    "total"),
        "sub_opened":    _sum(subs,    "opened"),
        "sub_clicked":   _sum(subs,    "clicked"),
        "sub_replied":   _sum(subs,    "replied"),
    }

    # ── Positive leads (with full conversation messages) ──────────────
    POSITIVE = {"interested", "meeting booked", "positive",
                "meeting request", "will buy", "warm", "demo request"}
    pos_leads = []
    for p in parents:
        for l in p.get("leads", []):
            if (l.get("category") or "").lower() in POSITIVE:
                pos_leads.append(_slim_lead(l, p.get("client", "")))
    stats["positive"] = len(pos_leads)

    # ── LinkedIn — filter campaigns/convos/leads by client keyword ────
    li_campaigns, li_convos, li_leads = [], [], []
    try:
        accounts, lcamps, lrecent, lconvos = get_li_data()
        acc_map = {a["id"]: a.get("full_name", a["id"]) for a in accounts}
        match = [c for c in lcamps if _matches_client(c.get("name"), keywords)]
        camp_ids = {c["id"] for c in match}
        camp_name = {c["id"]: c.get("name", "") for c in match}
        for c in match:
            li_campaigns.append({
                "id": c["id"], "name": c.get("name", ""),
                "state": c.get("state", ""), "type": c.get("type", ""),
                "targets": c.get("targets", 0),
                "owners": [acc_map.get(o, o) for o in _owners_list(c.get("owners"))],
                "completion_pct": c.get("completion_pct", 0),
                "accepted": c.get("accepted", 0),
                "replies":  c.get("replies",  0),
            })
        for ev in lrecent:
            if ev.get("campaign_id") in camp_ids:
                li_leads.append({
                    "transition": ev.get("transition", ""),
                    "lead":       ev.get("lead_name", ev.get("lead", "")),
                    "occupation": ev.get("occupation", ""),
                    "campaign":   camp_name.get(ev.get("campaign_id"), ""),
                    "account":    acc_map.get(ev.get("account_id"), ev.get("account_id", "")),
                    "date":       (ev.get("timestamp") or "")[:10],
                })
        for cv in lconvos:
            if cv.get("campaign_id") in camp_ids or _matches_client(cv.get("campaign_name"), keywords):
                li_convos.append({
                    "lead":       cv.get("lead_name", ""),
                    "occupation": cv.get("occupation", ""),
                    "campaign":   cv.get("campaign_name", ""),
                    "account":    acc_map.get(cv.get("account_id"), cv.get("account_id", "")),
                    "msgs":       cv.get("msg_count", 0),
                    "unread":     cv.get("unread", 0),
                    "last":       (cv.get("last_message_time") or "")[:19],
                })
    except Exception as e:
        log.warning("Client LinkedIn fetch error: %s", e)

    # ── Zoho — pull this client's by_company entry if present ─────────
    zoho = {"last_sync": None, "created": 0, "updated": 0, "total": 0}
    try:
        zpath = os.path.join(os.path.dirname(__file__), "reports", "zoho_sync_stats.json")
        with open(zpath) as f:
            zdata = json.load(f)
        zoho["last_sync"] = zdata.get("last_sync")
        bc = (zdata.get("by_company") or {}).get(name) or {}
        zoho["created"] = bc.get("created", 0)
        zoho["updated"] = bc.get("updated", 0)
        zoho["total"]   = zoho["created"] + zoho["updated"]
    except FileNotFoundError:
        pass
    except Exception as e:
        log.warning("Client Zoho fetch error: %s", e)

    return jsonify({
        "client":          name,
        "slug":            slug,
        "loading":         email.get("loading") or _email_fetching,
        "stats":           stats,
        "parent_campaigns": parents,
        "subsequences":    subs,
        "positive_leads":  pos_leads,
        "linkedin": {
            "campaigns":     li_campaigns,
            "lead_events":   li_leads[:200],
            "conversations": li_convos[:200],
        },
        "zoho": zoho,
    })


@app.route("/client/<slug>")
def client_dashboard(slug):
    client = CLIENTS_BY_SLUG.get(slug)
    if not client:
        return f"Unknown client: {slug}", 404
    return render_template_string(CLIENT_DASHBOARD_HTML,
                                  client=client, clients=CLIENTS)


@app.route("/")
def index():
    return render_template_string(DASHBOARD_HTML, clients=CLIENTS)


# ── HTML Dashboard ────────────────────────────────────────────────────

DASHBOARD_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Unified Analytics Dashboard</title>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:#f0f2f5;color:#1e293b;min-height:100vh}

/* ── Topbar ── */
.topbar{background:linear-gradient(135deg,#1e1b4b 0%,#3730a3 55%,#4f46e5 100%);padding:0 28px;display:flex;align-items:center;justify-content:space-between;height:56px;position:sticky;top:0;z-index:200;box-shadow:0 4px 20px rgba(67,56,202,.4)}
.logo{font-size:16px;font-weight:800;color:#fff;display:flex;align-items:center;gap:8px}
.logo-mark{width:28px;height:28px;border-radius:8px;background:linear-gradient(135deg,#818cf8,#a78bfa);display:flex;align-items:center;justify-content:center}
.topbar-right{display:flex;align-items:center;gap:10px}
.ts-badge{font-size:11px;color:rgba(255,255,255,.5)}
.cached-pill{font-size:10px;color:#fbbf24;background:rgba(251,191,36,.15);border:1px solid rgba(251,191,36,.3);padding:2px 8px;border-radius:20px;font-weight:700}
.refresh-btn{background:rgba(255,255,255,.12);color:#fff;border:1px solid rgba(255,255,255,.22);padding:7px 16px;border-radius:8px;font-size:12px;font-weight:600;cursor:pointer;display:flex;align-items:center;gap:5px;transition:all .2s}
.refresh-btn:hover{background:rgba(255,255,255,.22)}
.refresh-btn:disabled{opacity:.4;cursor:not-allowed}

/* ── Section wrappers ── */
.section-block{background:rgba(255,255,255,.0);padding:20px 28px 4px}
.section-label{display:flex;align-items:center;gap:10px;margin-bottom:16px}
.section-label h2{font-size:17px;font-weight:800;color:#0f172a}
.section-label .src-pill{font-size:11px;font-weight:700;padding:3px 10px;border-radius:20px}
.li-pill{background:#dbeafe;color:#1d4ed8}
.sm-pill{background:#fef3c7;color:#92400e}
.section-divider{height:1px;background:linear-gradient(90deg,#e2e8f0,transparent);margin:8px 28px 0}

/* ── Filter bars ── */
.filter-bar{background:rgba(255,255,255,.9);backdrop-filter:blur(12px);border:1px solid #e2e8f0;border-radius:12px;padding:12px 18px;margin-bottom:16px;display:flex;align-items:center;gap:14px;flex-wrap:wrap;box-shadow:0 2px 12px rgba(0,0,0,.05)}
.fl{font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:1px;white-space:nowrap}
.fdiv{width:1px;height:28px;background:#e2e8f0;flex-shrink:0}
.fsel,.fdate{border:1.5px solid #e2e8f0;border-radius:7px;padding:6px 10px;font-size:12px;color:#1e293b;background:#f8fafc;font-family:inherit;outline:none;cursor:pointer}
.fsel:focus,.fdate:focus{border-color:#6366f1;background:#eef2ff}
.fbtn{padding:6px 14px;background:#4f46e5;color:#fff;border:none;border-radius:7px;font-size:12px;font-weight:700;cursor:pointer;font-family:inherit}
.fbtn:hover{background:#4338ca}
.fbtn-reset{background:#f1f5f9;color:#64748b;border:1px solid #e2e8f0}
.fbtn-reset:hover{background:#e2e8f0}
.fstatus{font-size:11px;color:#6366f1;font-weight:600}

/* ── Stat cards ── */
.stats-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:18px}
.stats-grid-8{display:grid;grid-template-columns:repeat(8,1fr);gap:12px;margin-bottom:18px}
@media(max-width:1100px){.stats-grid-8{grid-template-columns:repeat(4,1fr)}}
@media(max-width:900px){.stats-grid{grid-template-columns:repeat(2,1fr)}}
.stat-card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:16px;box-shadow:0 1px 6px rgba(0,0,0,.05);cursor:pointer;transition:transform .15s,box-shadow .15s}
.stat-card:hover{transform:translateY(-2px);box-shadow:0 6px 20px rgba(0,0,0,.1)}
.stat-icon{width:34px;height:34px;border-radius:10px;display:flex;align-items:center;justify-content:center;margin-bottom:10px}
.stat-val{font-size:26px;font-weight:900;line-height:1;letter-spacing:-1px}
.stat-lbl{font-size:11px;color:#64748b;font-weight:600;margin-top:3px}
.stat-sub{font-size:10px;font-weight:700;margin-top:4px}

/* Colour themes */
.c-blue .stat-val,.c-blue .stat-sub{color:#2563eb} .c-blue .stat-icon{background:#dbeafe}
.c-green .stat-val,.c-green .stat-sub{color:#16a34a} .c-green .stat-icon{background:#dcfce7}
.c-purple .stat-val,.c-purple .stat-sub{color:#7c3aed} .c-purple .stat-icon{background:#ede9fe}
.c-orange .stat-val,.c-orange .stat-sub{color:#ea580c} .c-orange .stat-icon{background:#fff7ed}
.c-cyan .stat-val,.c-cyan .stat-sub{color:#0891b2} .c-cyan .stat-icon{background:#cffafe}
.c-amber .stat-val,.c-amber .stat-sub{color:#d97706} .c-amber .stat-icon{background:#fef3c7}
.c-rose .stat-val,.c-rose .stat-sub{color:#e11d48} .c-rose .stat-icon{background:#fecdd3}
.c-slate .stat-val,.c-slate .stat-sub{color:#475569} .c-slate .stat-icon{background:#f1f5f9}
.c-emerald .stat-val,.c-emerald .stat-sub{color:#059669} .c-emerald .stat-icon{background:#d1fae5}

/* ── Cards ── */
.card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;box-shadow:0 1px 6px rgba(0,0,0,.05);margin-bottom:16px;overflow:hidden}
.card-head{padding:13px 18px;border-bottom:1px solid #f1f5f9;display:flex;align-items:center;justify-content:space-between}
.card-head-left{display:flex;align-items:center;gap:8px}
.card-head-icon{width:30px;height:30px;border-radius:8px;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.card-head h3{font-size:13px;font-weight:800;color:#1e293b}
.count-pill{font-size:11px;color:#64748b;background:#f1f5f9;padding:2px 10px;border-radius:20px;font-weight:700}

/* ── Tables ── */
table{width:100%;border-collapse:collapse;font-size:12px}
th{padding:9px 14px;text-align:left;font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:.7px;border-bottom:1px solid #f1f5f9;background:#fafbfc;white-space:nowrap}
td{padding:10px 14px;border-bottom:1px solid #f8fafc;color:#374151;vertical-align:middle}
tr:last-child td{border-bottom:none}
tbody tr:hover td{background:#fafbff}
.clickable{cursor:pointer;color:#2563eb;font-weight:600;text-decoration:underline dotted}
.clickable:hover{color:#1d4ed8}
.rate{color:#94a3b8;font-size:10px}
.no-data{text-align:center;padding:28px;color:#94a3b8;font-size:13px}

/* ── Badges ── */
.badge{display:inline-flex;align-items:center;padding:2px 8px;border-radius:20px;font-size:10px;font-weight:800;white-space:nowrap}
.b-active{background:#dcfce7;color:#15803d;border:1px solid #86efac}
.b-done{background:#ede9fe;color:#6d28d9;border:1px solid #c4b5fd}
.b-init{background:#fef3c7;color:#92400e;border:1px solid #fcd34d}
.b-paused{background:#fee2e2;color:#b91c1c;border:1px solid #fca5a5}
.b-accepted{background:#dcfce7;color:#15803d;border:1px solid #86efac}
.b-reply{background:#dbeafe;color:#1d4ed8;border:1px solid #93c5fd}

/* ── Charts ── */
.charts-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(360px,1fr));gap:14px;margin-bottom:16px}
.chart-card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:14px;box-shadow:0 1px 6px rgba(0,0,0,.05)}
.chart-card h4{font-size:11px;font-weight:700;color:#64748b;margin-bottom:10px}
.chart-wrap{position:relative;height:220px}

/* ── Conversations ── */
.conv-list{padding:10px 14px}
.conv-item{border:1px solid #e2e8f0;border-radius:10px;margin-bottom:8px;overflow:hidden}
.conv-header{padding:11px 14px;background:#fafafa;display:flex;align-items:center;gap:10px;cursor:pointer;transition:background .1s}
.conv-header:hover{background:#f1f5f9}
.av{width:36px;height:36px;border-radius:50%;background:linear-gradient(135deg,#dbeafe,#ede9fe);color:#4338ca;font-size:13px;font-weight:800;display:flex;align-items:center;justify-content:center;flex-shrink:0;overflow:hidden}
.av img{width:36px;height:36px;border-radius:50%;object-fit:cover}
.conv-info{flex:1;min-width:0}
.conv-name{font-size:12px;font-weight:700;color:#1e293b}
.conv-occ{font-size:10px;color:#94a3b8;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;margin-top:1px}
.conv-right{display:flex;align-items:center;gap:6px;flex-shrink:0}
.unread-pill{background:linear-gradient(135deg,#ef4444,#dc2626);color:#fff;border-radius:20px;padding:1px 7px;font-size:10px;font-weight:700}
.conv-body{display:none;padding:14px;background:#fff;border-top:1px solid #f1f5f9;max-height:320px;overflow-y:auto}
.conv-body.open{display:block}
.msg-row{display:flex;gap:7px;margin-bottom:10px;align-items:flex-start}
.msg-row.sent{flex-direction:row-reverse}
.msg-av{width:26px;height:26px;border-radius:50%;background:#f1f5f9;color:#64748b;font-size:10px;font-weight:700;display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:2px}
.msg-bubble{max-width:70%;padding:8px 11px;border-radius:12px;font-size:11.5px;line-height:1.5}
.msg-row.received .msg-bubble{background:#f1f5f9;color:#1e293b}
.msg-row.sent .msg-bubble{background:linear-gradient(135deg,#eff6ff,#dbeafe);color:#1d4ed8}
.msg-meta{font-size:10px;color:#94a3b8;margin-bottom:3px}
.auto-tag{color:#7c3aed;font-weight:700}
.chev{font-size:9px;color:#94a3b8;transition:transform .2s}

/* ── Zoho widget ── */
.zoho-chips{display:flex;gap:10px;flex-wrap:wrap;padding:14px 18px}
.zoho-chip{border-radius:10px;padding:10px 14px;cursor:pointer;transition:opacity .15s;min-width:90px}
.zoho-chip:hover{opacity:.8}
.zoho-chip-val{font-size:22px;font-weight:900;line-height:1}
.zoho-chip-lbl{font-size:10px;opacity:.7;margin-top:1px;font-weight:600}
.zoho-sync-btn{background:linear-gradient(135deg,#0ea5e9,#0369a1);color:#fff;border:none;padding:6px 14px;border-radius:8px;font-size:11px;font-weight:700;cursor:pointer;display:flex;align-items:center;gap:5px;transition:opacity .2s}
.zoho-sync-btn:disabled{opacity:.45;cursor:not-allowed}

/* ── Email modal ── */
.modal-overlay{display:none;position:fixed;inset:0;background:rgba(15,23,42,.5);z-index:500;align-items:center;justify-content:center}
.modal-overlay.open{display:flex}
.modal{background:#fff;border-radius:14px;width:92vw;max-width:1080px;height:80vh;display:flex;flex-direction:column;overflow:hidden;box-shadow:0 20px 60px rgba(0,0,0,.2)}
.modal-header{padding:13px 18px;border-bottom:1px solid #e2e8f0;display:flex;justify-content:space-between;align-items:center;flex-shrink:0;background:#f8fafc}
.modal-header h3{font-size:13px;font-weight:800;color:#0f172a}
.modal-close{background:none;border:none;color:#94a3b8;font-size:20px;cursor:pointer;line-height:1}
.modal-body{display:flex;flex:1;overflow:hidden}
.lead-list{width:270px;border-right:1px solid #e2e8f0;overflow-y:auto;flex-shrink:0;background:#f8fafc}
.lead-item{padding:10px 13px;border-bottom:1px solid #f1f5f9;cursor:pointer}
.lead-item:hover{background:#f1f5f9}
.lead-item.active{background:#eff6ff;border-left:3px solid #2563eb}
.lead-name{font-size:12px;font-weight:600;color:#1e293b}
.lead-email{font-size:10px;color:#64748b;margin-top:1px;word-break:break-all}
.lead-meta{font-size:9px;color:#94a3b8;margin-top:3px}
.msg-panel{flex:1;overflow-y:auto;padding:16px;background:#fff}
.msg-panel .placeholder{color:#94a3b8;font-size:12px;text-align:center;padding-top:48px}
.sm-msg-bubble{background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;padding:11px;margin-bottom:11px}
.sm-msg-bubble.inbound{background:#fffbeb;border-color:#fde68a;border-left:3px solid #f59e0b}
.sm-msg-meta{display:flex;justify-content:space-between;margin-bottom:5px;flex-wrap:wrap;gap:3px}
.sm-msg-seq{font-size:10px;color:#2563eb;font-weight:700}
.sm-msg-date{font-size:10px;color:#94a3b8}
.sm-msg-subj{font-size:12px;font-weight:600;color:#1e293b;margin-bottom:5px}
.sm-msg-body{font-size:11px;color:#475569;line-height:1.6;white-space:pre-wrap;word-break:break-word}
.sm-tags{display:flex;gap:5px;margin-top:6px;flex-wrap:wrap}
.sm-tag{display:inline-block;padding:1px 6px;border-radius:3px;font-size:9px;font-weight:700}
.tag-open{background:#dcfce7;color:#16a34a}
.tag-click{background:#dbeafe;color:#2563eb}
.tag-reply{background:#fef3c7;color:#d97706}
.dir-out{background:#dbeafe;color:#1d4ed8;font-size:9px;font-weight:700;padding:1px 5px;border-radius:3px;margin-right:4px}
.dir-in{background:#fef3c7;color:#b45309;font-size:9px;font-weight:700;padding:1px 5px;border-radius:3px;margin-right:4px}

/* ── LinkedIn drawer ── */
.drawer-backdrop{position:fixed;inset:0;background:rgba(15,23,42,.35);backdrop-filter:blur(3px);z-index:500;opacity:0;pointer-events:none;transition:opacity .25s}
.drawer-backdrop.open{opacity:1;pointer-events:all}
.drawer{position:fixed;top:0;right:0;bottom:0;width:660px;max-width:95vw;background:#fff;z-index:501;transform:translateX(100%);transition:transform .28s cubic-bezier(.4,0,.2,1);display:flex;flex-direction:column;box-shadow:-8px 0 40px rgba(0,0,0,.12)}
.drawer.open{transform:translateX(0)}
.drawer-head{padding:18px 22px;border-bottom:1px solid #f1f5f9;display:flex;align-items:center;gap:10px;flex-shrink:0}
.drawer-head-icon{width:34px;height:34px;border-radius:10px;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.drawer-title{font-size:14px;font-weight:800;color:#1e293b;flex:1}
.drawer-count{font-size:11px;color:#64748b;background:#f1f5f9;padding:2px 10px;border-radius:20px;font-weight:700}
.drawer-close{width:30px;height:30px;border-radius:7px;border:none;background:#f1f5f9;color:#64748b;font-size:16px;cursor:pointer;display:flex;align-items:center;justify-content:center;transition:background .15s}
.drawer-close:hover{background:#e2e8f0;color:#1e293b}
.drawer-body{flex:1;overflow-y:auto}
.drawer th{position:sticky;top:0;z-index:1;background:#fafbfc}
.drawer td{padding:10px 14px}
.lead-cell{display:flex;align-items:center;gap:8px}
.lead-av{width:30px;height:30px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:700;flex-shrink:0;overflow:hidden}
.lead-av img{width:30px;height:30px;border-radius:50%;object-fit:cover}

/* ── Loading ── */
#loading{position:fixed;inset:0;background:linear-gradient(135deg,rgba(238,242,255,.97),rgba(245,243,255,.97));display:flex;flex-direction:column;align-items:center;justify-content:center;z-index:999;gap:12px}
.spinner{width:44px;height:44px;border:3px solid #e2e8f0;border-top-color:#4f46e5;border-right-color:#a855f7;border-radius:50%;animation:spin .7s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}
.load-txt{font-size:15px;color:#4338ca;font-weight:800}
.load-sub{font-size:11px;color:#94a3b8}

/* ── Error ── */
.err-bar{background:#fef2f2;border:1px solid #fecaca;color:#dc2626;padding:10px 16px;border-radius:8px;margin-bottom:14px;font-size:12px;display:none;font-weight:600}

/* ── Funnel bar (email) ── */
.funnel-row{display:flex;align-items:center;gap:8px;margin-bottom:7px}
.funnel-lbl{width:100px;font-size:11px;color:#64748b;text-align:right;flex-shrink:0}
.funnel-bar-wrap{flex:1;background:#f1f5f9;border-radius:4px;height:20px;overflow:hidden}
.funnel-bar{height:100%;border-radius:4px;display:flex;align-items:center;padding-left:7px;font-size:10px;font-weight:700;color:#fff;transition:width .4s}
.funnel-val{width:70px;font-size:11px;color:#475569;flex-shrink:0}

@keyframes fadeUp{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:translateY(0)}}
.fade-in{animation:fadeUp .3s ease-out forwards}

/* ── Pagination ── */
.pg-wrap{display:flex;align-items:center;justify-content:flex-end;gap:8px;padding:10px 14px;border-top:1px solid #f1f5f9;background:#fafbfc}
.pg-btn{background:#f1f5f9;border:1px solid #e2e8f0;color:#374151;padding:4px 12px;border-radius:6px;font-size:11px;font-weight:600;cursor:pointer;transition:all .15s;font-family:inherit}
.pg-btn:hover:not(:disabled){background:#e2e8f0;border-color:#cbd5e1}
.pg-btn:disabled{opacity:.38;cursor:not-allowed}
.pg-info{font-size:11px;color:#64748b;font-weight:600}

/* ── Breadcrumb ── */
.breadcrumb{display:flex;align-items:center;gap:6px;padding:6px 0 10px;font-size:12px;flex-wrap:wrap}
.bc-item{font-weight:600;cursor:pointer;color:#4f46e5;transition:color .15s}
.bc-item:hover{color:#3730a3;text-decoration:underline}
.bc-sep{color:#cbd5e1;font-size:11px}
.bc-cur{color:#1e293b;font-weight:700}

/* ── Topbar action buttons ── */
.act-btn{padding:5px 13px;border:none;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer;display:flex;align-items:center;gap:5px;transition:opacity .15s;font-family:inherit}
.act-btn:hover{opacity:.84}
.act-btn:disabled{opacity:.45;cursor:not-allowed}
.act-pdf{background:linear-gradient(135deg,#ef4444,#dc2626);color:#fff}
.act-docx{background:linear-gradient(135deg,#2563eb,#1d4ed8);color:#fff}
.act-share{background:linear-gradient(135deg,#059669,#047857);color:#fff}

/* ── LinkedIn Segment Stat Card ── */
.li-seg-filter{display:flex;align-items:center;gap:10px;padding:12px 16px;border-bottom:1px solid #f1f5f9;flex-wrap:wrap}
.li-seg-cards{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;padding:14px 16px}
@media(max-width:800px){.li-seg-cards{grid-template-columns:repeat(2,1fr)}}

/* ── Positive Leads section ── */
#pos-leads-section{margin-bottom:32px}

/* ── Share via Email modal ── */
.share-overlay{display:none;position:fixed;inset:0;background:rgba(15,23,42,.52);z-index:600;align-items:center;justify-content:center}
.share-overlay.open{display:flex}
.share-box{background:#fff;border-radius:14px;width:420px;max-width:95vw;padding:26px;box-shadow:0 24px 64px rgba(0,0,0,.22)}
.share-box h3{font-size:15px;font-weight:800;color:#0f172a;margin-bottom:6px}
.share-box p{font-size:12px;color:#64748b;margin-bottom:16px}
.share-inp{width:100%;border:1.5px solid #e2e8f0;border-radius:8px;padding:9px 12px;font-size:13px;font-family:inherit;outline:none;margin-bottom:12px;box-sizing:border-box}
.share-inp:focus{border-color:#4f46e5;box-shadow:0 0 0 3px rgba(79,70,229,.1)}
.share-row{display:flex;gap:8px;justify-content:flex-end}
.share-send{background:linear-gradient(135deg,#4f46e5,#7c3aed);color:#fff;border:none;border-radius:8px;padding:8px 20px;font-size:13px;font-weight:700;cursor:pointer;font-family:inherit;transition:opacity .15s}
.share-send:disabled{opacity:.5;cursor:not-allowed}
.share-cancel{background:#f1f5f9;color:#64748b;border:1px solid #e2e8f0;border-radius:8px;padding:8px 16px;font-size:12px;font-weight:600;cursor:pointer;font-family:inherit}
.share-msg{margin-top:10px;font-size:12px;font-weight:600;padding:8px 12px;border-radius:6px;display:none}
.share-ok{background:#dcfce7;color:#15803d;display:block}
.share-err{background:#fee2e2;color:#dc2626;display:block}
</style>
</head>
<body>

<div id="loading">
  <div class="spinner"></div>
  <div class="load-txt">Unified Analytics</div>
  <div class="load-sub" id="load-sub-txt">Loading LinkedIn data…</div>
</div>

<!-- Topbar -->
<div class="topbar">
  <div class="logo">
    <div class="logo-mark">
      <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#fff" stroke-width="2.5" stroke-linecap="round"><polyline points="22 12 18 12 15 21 9 3 6 12 2 12"/></svg>
    </div>
    Unified Analytics Dashboard
  </div>
  <div class="topbar-right">
    <span class="ts-badge" id="li-ts"></span>
    <span class="ts-badge" id="sm-ts" style="margin-left:8px"></span>
    <span class="cached-pill" id="li-cached" style="display:none">LI Cached</span>
    <span class="cached-pill" id="sm-cached" style="display:none">Email Cached</span>
    <button class="act-btn act-pdf" id="pdf-btn" onclick="downloadPDF()">&#x21E9; PDF</button>
    <button class="act-btn act-docx" id="docx-btn" onclick="downloadDOCX()">&#x21E9; DOCX</button>
    <button class="act-btn act-share" onclick="openShareModal()">&#x2709; Share</button>
    <button class="refresh-btn" id="refresh-btn" onclick="refreshAll()">
      <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg>
      Refresh All
    </button>
  </div>
</div>

<!-- Client tab navigation -->
<div style="background:#fff;border-bottom:1px solid #e2e8f0;padding:10px 28px;display:flex;gap:6px;flex-wrap:wrap;align-items:center;position:sticky;top:56px;z-index:150">
  <span style="font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:1px;margin-right:6px">View:</span>
  <a href="/" style="padding:6px 14px;border-radius:7px;font-size:12px;font-weight:700;text-decoration:none;background:#4f46e5;color:#fff">All Clients</a>
  {% for c in clients %}
  <a href="/client/{{c.slug}}" style="padding:6px 14px;border-radius:7px;font-size:12px;font-weight:700;text-decoration:none;background:#f1f5f9;color:#475569;border:1px solid #e2e8f0">{{c.name}}</a>
  {% endfor %}
</div>

<div style="max-width:1500px;margin:0 auto;padding:20px 24px">

  <div class="err-bar" id="li-err"></div>
  <div class="err-bar" id="sm-err"></div>

  <!-- ═══════════════════════════════════════════════════════════════ -->
  <!--  LINKEDIN SECTION                                              -->
  <!-- ═══════════════════════════════════════════════════════════════ -->
  <div class="section-block" style="padding:0 0 4px">
    <div class="section-label">
      <h2>🔗 LinkedIn Overview</h2>
      <span class="src-pill li-pill">Aimfox</span>
    </div>

    <!-- LinkedIn Filter -->
    <div class="filter-bar" id="li-filter-bar">
      <span class="fl">Account</span>
      <div style="position:relative;min-width:220px">
        <select class="fsel" id="li-acc" onchange="renderLI()" style="width:100%">
          <option value="all">All Accounts</option>
        </select>
      </div>
      <div class="fdiv"></div>
      <span class="fl">Date Range</span>
      <input type="date" class="fdate" id="li-from" onchange="renderLI()">
      <span style="font-size:11px;color:#94a3b8">to</span>
      <input type="date" class="fdate" id="li-to" onchange="renderLI()">
      <button class="fbtn" onclick="setLIToday()">Today</button>
      <button class="fbtn" onclick="setLIWeek()">This Week</button>
      <button class="fbtn fbtn-reset" onclick="clearLIDate()">Clear</button>
    </div>

    <!-- LinkedIn KPI Cards -->
    <div class="stats-grid" id="li-stats"></div>

    <!-- Zoho CRM Sync -->
    <div class="card">
      <div class="card-head">
        <div class="card-head-left">
          <div class="card-head-icon" style="background:#e0f2fe">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#0369a1" stroke-width="2" stroke-linecap="round"><ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M21 12c0 1.66-4 3-9 3s-9-1.34-9-3"/><path d="M3 5v14c0 1.66 4 3 9 3s9-1.34 9-3V5"/></svg>
          </div>
          <h3>Zoho CRM Sync</h3>
        </div>
        <div style="display:flex;align-items:center;gap:8px">
          <span id="zoho-ts" style="font-size:10px;color:#94a3b8;font-weight:600"></span>
          <button class="zoho-sync-btn" id="zoho-sync-btn" onclick="triggerZohoSync()">
            <svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg>
            Sync Now
          </button>
        </div>
      </div>
      <div class="zoho-chips" id="zoho-body"><span style="color:#94a3b8;font-size:12px">Loading…</span></div>
    </div>

    <!-- Conversations -->
    <div class="card">
      <div class="card-head">
        <div class="card-head-left">
          <div class="card-head-icon" style="background:#ede9fe">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#4338ca" stroke-width="2" stroke-linecap="round"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>
          </div>
          <h3>Conversations</h3>
        </div>
        <span class="count-pill" id="li-conv-count"></span>
      </div>
      <div class="conv-list" id="li-conv-list"></div>
    </div>

    <!-- LinkedIn Campaigns -->
    <div class="card">
      <div class="card-head">
        <div class="card-head-left">
          <div class="card-head-icon" style="background:#dcfce7">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#15803d" stroke-width="2" stroke-linecap="round"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg>
          </div>
          <h3>LinkedIn Campaigns</h3>
        </div>
        <span class="count-pill" id="li-camp-count"></span>
      </div>
      <table><thead><tr>
        <th>Campaign</th><th>State</th><th>Type</th><th>Targets</th><th>Accepted</th><th>Replies</th><th>Account</th><th>Created</th>
      </tr></thead><tbody id="li-camp-body"></tbody></table>
    </div>

    <!-- Recent LinkedIn Leads -->
    <div class="card">
      <div class="card-head">
        <div class="card-head-left">
          <div class="card-head-icon" style="background:#fef3c7">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#92400e" stroke-width="2" stroke-linecap="round"><polygon points="13 2 3 14 12 14 11 22 21 10 12 10 13 2"/></svg>
          </div>
          <h3>Recent Lead Events</h3>
        </div>
        <span class="count-pill" id="li-leads-count"></span>
      </div>
      <table><thead><tr>
        <th>Event</th><th>Lead</th><th>Occupation</th><th>Campaign</th><th>Account</th><th>Date</th>
      </tr></thead><tbody id="li-leads-body"></tbody></table>
    </div>
  </div>

  <div class="section-divider" style="margin:4px 0 20px"></div>

  <!-- ═══════════════════════════════════════════════════════════════ -->
  <!--  EMAIL SECTION (Smartlead)                                     -->
  <!-- ═══════════════════════════════════════════════════════════════ -->
  <div class="section-block" style="padding:0">
    <div class="section-label">
      <h2>📧 Email Analytics</h2>
      <span class="src-pill sm-pill">Smartlead</span>
      <span style="font-size:11px;color:#94a3b8" id="sm-loading-hint">Loading email data in background…</span>
    </div>

    <!-- Email Filter -->
    <div class="filter-bar">
      <span class="fl">Client</span>
      <select class="fsel" id="sm-client" onchange="applyEmailFilters()">
        <option value="ALL">All Clients</option>
      </select>
      <div class="fdiv"></div>
      <span class="fl">Date Range</span>
      <input type="date" class="fdate" id="sm-from">
      <span style="font-size:11px;color:#94a3b8">to</span>
      <input type="date" class="fdate" id="sm-to">
      <button class="fbtn" onclick="applyEmailFilters()">Apply</button>
      <button class="fbtn fbtn-reset" onclick="resetEmailFilters()">Reset</button>
      <span class="fstatus" id="sm-filter-status"></span>
    </div>

    <!-- Breadcrumb navigation -->
    <div class="breadcrumb" id="sm-breadcrumb">
      <span class="bc-cur">Home</span>
    </div>

    <!-- ── Combined Overview ── -->
    <div style="font-size:13px;font-weight:700;color:#475569;margin-bottom:10px;display:flex;align-items:center;gap:8px">
      🔗 Combined Overview <span style="font-size:11px;font-weight:400;color:#94a3b8">— full pipeline across campaigns & subsequences</span>
    </div>
    <div class="stats-grid-8" id="sm-combined-cards"></div>

    <div class="card" style="margin-bottom:16px">
      <div class="card-head"><div class="card-head-left"><h3>Full Pipeline by Client</h3></div></div>
      <div style="padding:14px"><div class="chart-wrap" style="height:260px"><canvas id="sm-pipeline-chart"></canvas></div></div>
    </div>

    <div class="card" style="margin-bottom:20px">
      <div class="card-head"><div class="card-head-left"><h3>Combined Table</h3></div></div>
      <div id="sm-combined-table"></div>
    </div>

    <!-- ── LinkedIn Segment Stat Card ── -->
    <div class="card" style="margin-bottom:20px">
      <div class="card-head">
        <div class="card-head-left">
          <div class="card-head-icon" style="background:#dbeafe">
            <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#1d4ed8" stroke-width="2" stroke-linecap="round"><path d="M16 8a6 6 0 0 1 6 6v7h-4v-7a2 2 0 0 0-2-2 2 2 0 0 0-2 2v7h-4v-7a6 6 0 0 1 6-6z"/><rect x="2" y="9" width="4" height="12"/><circle cx="4" cy="4" r="2"/></svg>
          </div>
          <h3>LinkedIn Stats — Segment View</h3>
        </div>
      </div>
      <div class="li-seg-filter">
        <span class="fl">Segment</span>
        <select class="fsel" id="li-stat-segment" onchange="loadLIStatCard()">
          <option value="all">All Accounts</option>
        </select>
        <div class="fdiv"></div>
        <span class="fl">Sub-segment</span>
        <select class="fsel" id="li-stat-sub" onchange="loadLIStatCard()">
          <option value="all">All Types</option>
        </select>
        <span id="li-stat-loading" style="font-size:11px;color:#6366f1;font-weight:600"></span>
        <span style="font-size:10px;color:#94a3b8;margin-left:4px">— auto-syncs with Client filter above</span>
      </div>
      <div class="li-seg-cards" id="li-stat-cards">
        <div style="color:#94a3b8;font-size:12px;padding:4px">Loading…</div>
      </div>
    </div>

    <!-- ── Main Campaign Analytics ── -->
    <div style="font-size:13px;font-weight:700;color:#475569;margin-bottom:10px;display:flex;align-items:center;gap:8px">
      📊 Main Campaign Analytics <span style="font-size:11px;font-weight:400;color:#94a3b8">— parent campaigns</span>
    </div>
    <div class="stats-grid" id="sm-main-cards"></div>
    <div class="charts-grid" style="margin-bottom:16px">
      <div class="chart-card"><h4>Opened Leads by Client</h4><div class="chart-wrap"><canvas id="sm-opened-chart"></canvas></div></div>
      <div class="chart-card"><h4>Added to Sub by Client</h4><div class="chart-wrap"><canvas id="sm-sub-chart"></canvas></div></div>
    </div>
    <div class="card" style="margin-bottom:20px">
      <div class="card-head"><div class="card-head-left"><h3>Campaign Details</h3></div><span class="count-pill" id="sm-camp-count"></span></div>
      <div id="sm-main-table"></div>
    </div>

    <!-- ── Subsequence Analytics ── -->
    <div style="font-size:13px;font-weight:700;color:#475569;margin-bottom:10px;display:flex;align-items:center;gap:8px">
      🔁 Subsequence Analytics <span style="font-size:11px;font-weight:400;color:#94a3b8">— double-click metric to see leads</span>
    </div>
    <div class="stats-grid" id="sm-sub-cards"></div>
    <div class="charts-grid" style="margin-bottom:16px">
      <div class="chart-card"><h4>Open Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="sm-open-rate-chart"></canvas></div></div>
      <div class="chart-card"><h4>Click Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="sm-click-rate-chart"></canvas></div></div>
      <div class="chart-card"><h4>Reply Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="sm-reply-rate-chart"></canvas></div></div>
      <div class="chart-card"><h4>Total Leads by Subsequence</h4><div class="chart-wrap"><canvas id="sm-leads-chart"></canvas></div></div>
    </div>
    <div class="card" style="margin-bottom:32px">
      <div class="card-head"><div class="card-head-left"><h3>Subsequence Details</h3></div><span class="count-pill" id="sm-sub-count"></span></div>
      <div id="sm-sub-table"></div>
    </div>

    <!-- ── Positive Leads ── -->
    <div id="pos-leads-section">
      <div style="font-size:13px;font-weight:700;color:#475569;margin-bottom:10px;display:flex;align-items:center;gap:8px">
        &#x2705; Positive Leads
        <span style="font-size:11px;font-weight:400;color:#94a3b8">— leads with positive engagement status</span>
        <span class="count-pill" id="pos-count">0</span>
      </div>
      <div class="card" style="margin-bottom:32px">
        <div class="card-head">
          <div class="card-head-left">
            <div class="card-head-icon" style="background:#dcfce7">
              <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#15803d" stroke-width="2" stroke-linecap="round"><polyline points="20 6 9 17 4 12"/></svg>
            </div>
            <h3>Positive Leads</h3>
          </div>
          <span class="count-pill" id="pos-table-count"></span>
        </div>
        <div id="pos-leads-table"></div>
      </div>
    </div>
  </div>
</div>

<!-- Share via Email Modal -->
<div class="share-overlay" id="share-overlay">
  <div class="share-box">
    <h3>&#x2709; Share Report via Email</h3>
    <p>A PDF of the current analytics will be generated and sent as an attachment.</p>
    <input class="share-inp" type="email" id="share-email-inp" placeholder="recipient@example.com">
    <div class="share-row">
      <button class="share-cancel" onclick="closeShareModal()">Cancel</button>
      <button class="share-send" id="share-send-btn" onclick="sendShareEmail()">Send Report</button>
    </div>
    <div class="share-msg" id="share-msg"></div>
  </div>
</div>

<!-- Email Lead Modal -->
<div class="modal-overlay" id="sm-modal">
  <div class="modal">
    <div class="modal-header">
      <h3 id="sm-modal-title">Lead Details</h3>
      <button class="modal-close" onclick="closeSMModal()">×</button>
    </div>
    <div class="modal-body">
      <div class="lead-list" id="sm-lead-list"></div>
      <div class="msg-panel" id="sm-msg-panel"><div class="placeholder">← Select a lead to view messages</div></div>
    </div>
  </div>
</div>

<!-- LinkedIn Detail Drawer -->
<div class="drawer-backdrop" id="li-drawer-backdrop" onclick="closeLIDrawer()"></div>
<div class="drawer" id="li-drawer">
  <div class="drawer-head">
    <div class="drawer-head-icon" id="li-drawer-icon"></div>
    <div class="drawer-title" id="li-drawer-title"></div>
    <span class="drawer-count" id="li-drawer-count"></span>
    <button class="drawer-close" onclick="closeLIDrawer()">×</button>
  </div>
  <div class="drawer-body" id="li-drawer-body"></div>
</div>

<script>
// ═══════════════════════════════════════════════════════════════════
//  STATE
// ═══════════════════════════════════════════════════════════════════
let LI = null;   // LinkedIn API data
let SM = null;   // Email API data
let _zoho = null;

// LinkedIn filters
let liAcc  = 'all', liFrom = '', liTo = '';

// Email filters
let smClient = 'ALL', smFrom = '', smTo = '';
let smParentFilt = [], smSubFilt = [];

// Chart instances
const charts = {};

// ── Paginator ────────────────────────────────────────────────────────
const pagers = {};
class Paginator {
  constructor(containerId, pageSize = 5) {
    this.id = containerId;
    this.ps = pageSize;
    this.page = 0;
    this.rows = [];
    this.hdr = '';
    this.rowFn = null;
  }
  init(rows, hdr, rowFn) {
    this.rows = rows; this.hdr = hdr; this.rowFn = rowFn; this.page = 0; this._render();
  }
  _total() { return Math.max(1, Math.ceil(this.rows.length / this.ps)); }
  _render() {
    const el = document.getElementById(this.id);
    if (!el) return;
    const start = this.page * this.ps;
    const slice = this.rows.slice(start, start + this.ps);
    const tp    = this._total();
    el.innerHTML = `<table><thead>${this.hdr}</thead><tbody>${
      slice.map((r,i) => this.rowFn(r, start + i)).join('')
    }</tbody></table><div class="pg-wrap">
      <span class="pg-info">Page ${this.page+1} of ${tp} &nbsp;(${this.rows.length} total)</span>
      <button class="pg-btn" onclick="pagers['${this.id}'].prev()" ${this.page===0?'disabled':''}>&#8592; Prev</button>
      <button class="pg-btn" onclick="pagers['${this.id}'].next()" ${this.page>=tp-1?'disabled':''}>Next &#8594;</button>
    </div>`;
  }
  next() { if (this.page < this._total()-1) { this.page++; this._render(); } }
  prev() { if (this.page > 0)              { this.page--; this._render(); } }
}
function mkPager(id, ps=5) { pagers[id] = new Paginator(id, ps); return pagers[id]; }

// ── Drill-down state ─────────────────────────────────────────────────
let drillClient = null;   // Level 1 (Segment)
let drillParent = null;   // Level 2 (Sub-segment / parent campaign raw_name)
let drillSub    = null;   // Level 3 (Sub-sub-segment / subsequence)

function setDrill(client, parent, sub) {
  drillClient = client; drillParent = parent; drillSub = sub;
  _renderBreadcrumb(); applyEmailFilters();
}

function _renderBreadcrumb() {
  const bc = document.getElementById('sm-breadcrumb');
  if (!bc) return;
  const parts = [`<span class="bc-item" onclick="setDrill(null,null,null)">Home</span>`];
  if (drillClient) {
    parts.push('<span class="bc-sep">&#x203A;</span>');
    if (!drillParent) parts.push(`<span class="bc-cur">${esc(drillClient)}</span>`);
    else parts.push(`<span class="bc-item" onclick="setDrill('${drillClient}',null,null)">${esc(drillClient)}</span>`);
  }
  if (drillParent) {
    parts.push('<span class="bc-sep">&#x203A;</span>');
    if (!drillSub) parts.push(`<span class="bc-cur">${esc(drillParent)}</span>`);
    else parts.push(`<span class="bc-item" onclick="setDrill('${drillClient}','${drillParent}',null)">${esc(drillParent)}</span>`);
  }
  if (drillSub) {
    parts.push('<span class="bc-sep">&#x203A;</span>');
    parts.push(`<span class="bc-cur">${esc(drillSub)}</span>`);
  }
  bc.innerHTML = parts.join('');
}

const POSITIVE_CATS = new Set(["interested","meeting booked","positive","meeting request","will buy","warm","demo request"]);
const SB = {ACTIVE:'b-active', DONE:'b-done', INIT:'b-init', PAUSED:'b-paused'};

function esc(s){ return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;') }
function fmt(n){ return (n||0).toLocaleString() }
function pct(n,t){ return t ? (n/t*100).toFixed(1)+'%' : '0%' }

// ═══════════════════════════════════════════════════════════════════
//  BOOT
// ═══════════════════════════════════════════════════════════════════
async function refreshAll() {
  const btn = document.getElementById('refresh-btn');
  btn.disabled = true;
  document.getElementById('loading').style.display = 'flex';
  document.getElementById('load-sub-txt').textContent = 'Refreshing…';
  await Promise.all([loadLI(), loadZoho()]);
  document.getElementById('loading').style.display = 'none';
  btn.disabled = false;
  loadEmail();  // email loads in background
}

window.onload = async () => {
  document.getElementById('load-sub-txt').textContent = 'Loading LinkedIn data…';
  await Promise.all([loadLI(), loadZoho()]);
  document.getElementById('loading').style.display = 'none';
  loadLIStatCard(true);   // populate LinkedIn segment stat card
  loadEmail();            // email loads in background
};

// ═══════════════════════════════════════════════════════════════════
//  LINKEDIN DATA
// ═══════════════════════════════════════════════════════════════════
async function loadLI() {
  const err = document.getElementById('li-err');
  err.style.display = 'none';
  try {
    const r = await fetch('/api/linkedin');
    if (!r.ok) throw new Error('HTTP ' + r.status);
    LI = await r.json();
    if (LI.error) throw new Error(LI.error);
    document.getElementById('li-ts').textContent = 'LinkedIn: ' + new Date().toLocaleTimeString();
    document.getElementById('li-cached').style.display = LI.cached ? 'inline-block' : 'none';
    populateLIAccounts();
    renderLI();
  } catch(e) {
    err.textContent = 'LinkedIn error: ' + e.message;
    err.style.display = 'block';
  }
}

function populateLIAccounts() {
  const sel = document.getElementById('li-acc');
  sel.innerHTML = '<option value="all">All Accounts</option>' +
    LI.accounts.map(a => `<option value="${a.id}">${esc(a.name)}</option>`).join('');
  sel.value = liAcc;
}

function curLIStats() {
  if (!liFrom && !liTo) {
    if (liAcc === 'all') return LI.overall;
    const acc = LI.accounts.find(a => String(a.id) === String(liAcc));
    return acc ? acc.stats : LI.overall;
  }
  // recompute from filtered events
  const fl = filtLILeads();
  const accepted = fl.filter(l => l.transition === 'accepted').length;
  const replies  = fl.filter(l => l.transition === 'reply').length;
  let messages_sent = 0, active_conversations = 0;
  for (const conv of filtLIConvos()) {
    let hasAct = false;
    for (const msg of conv.messages) {
      if (_liInRange(msg.date) && (msg.automated || msg.sender === conv.owner_name)) {
        messages_sent++; hasAct = true;
      }
    }
    if (hasAct) active_conversations++;
  }
  const base = liAcc === 'all' ? LI.overall : (LI.accounts.find(a=>String(a.id)===String(liAcc))||{stats:LI.overall}).stats;
  return { ...base, targets_sent: null, accepted, replies, messages_sent, active_conversations,
           accept_rate:'—', reply_rate: accepted>0 ? (replies/accepted*100).toFixed(1)+'%' : '—' };
}

function _liInRange(d) {
  if (!liFrom && !liTo) return true;
  if (liFrom && d < liFrom) return false;
  if (liTo   && d > liTo)   return false;
  return true;
}

function filtLILeads() {
  let leads = LI.recent_leads;
  if (liAcc !== 'all') leads = leads.filter(l => String(l.account_id) === String(liAcc));
  if (liFrom || liTo)  leads = leads.filter(l => _liInRange(l.timestamp));
  return leads;
}

function filtLIConvos() {
  let convos = liAcc === 'all' ? LI.conversations : LI.conversations.filter(c => String(c.owner) === String(liAcc));
  if (liFrom || liTo) convos = convos.filter(c => c.messages.some(m => _liInRange(m.date)));
  return convos;
}

function filtLICamps() {
  if (liAcc === 'all') return LI.campaigns;
  return LI.campaigns.filter(c => Array.isArray(c.owners) && c.owners.includes(liAcc));
}

function renderLI() {
  if (!LI) return;
  liAcc  = document.getElementById('li-acc').value;
  liFrom = document.getElementById('li-from').value;
  liTo   = document.getElementById('li-to').value;
  renderLIStats(); renderLIConvos(); renderLICamps(); renderLILeads();
}

function renderLIStats() {
  const s = curLIStats();
  const noDate = s.targets_sent === null;
  const tVal = noDate ? '<div style="font-size:18px;color:#94a3b8;font-weight:700">N/A</div>'
    : `<div class="stat-val">${fmt(s.targets_sent)}</div>`;
  document.getElementById('li-stats').innerHTML = `
    <div class="stat-card c-blue" onclick="showLIDetail('requests')">
      <div class="stat-icon"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#2563eb" stroke-width="2.5" stroke-linecap="round"><line x1="22" y1="2" x2="11" y2="13"/><polygon points="22 2 15 22 11 13 2 9 22 2"/></svg></div>
      ${tVal}<div class="stat-lbl">Requests Sent</div>
      <div class="stat-sub">${noDate ? 'Per-day N/A' : s.total_campaigns+' campaigns'}</div>
    </div>
    <div class="stat-card c-green" onclick="showLIDetail('accepted')">
      <div class="stat-icon"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#16a34a" stroke-width="2.5" stroke-linecap="round"><path d="M17 21v-2a4 4 0 0 0-4-4H5a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M23 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/></svg></div>
      <div class="stat-val">${fmt(s.accepted)}</div><div class="stat-lbl">Connections Accepted</div>
      <div class="stat-sub">Rate: ${s.accept_rate}</div>
    </div>
    <div class="stat-card c-purple" onclick="showLIDetail('messages')">
      <div class="stat-icon"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#7c3aed" stroke-width="2.5" stroke-linecap="round"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg></div>
      <div class="stat-val">${fmt(s.messages_sent)}</div><div class="stat-lbl">Messages Sent</div>
      <div class="stat-sub">${s.active_conversations} active convos</div>
    </div>
    <div class="stat-card c-orange" onclick="showLIDetail('replies')">
      <div class="stat-icon"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#ea580c" stroke-width="2.5" stroke-linecap="round"><polyline points="9 17 4 12 9 7"/><path d="M20 18v-2a4 4 0 0 0-4-4H4"/></svg></div>
      <div class="stat-val">${fmt(s.replies)}</div><div class="stat-lbl">Replies Received</div>
      <div class="stat-sub">Rate: ${s.reply_rate}</div>
    </div>`;
}

function renderLIConvos() {
  const convos = filtLIConvos();
  document.getElementById('li-conv-count').textContent = convos.length;
  document.getElementById('li-conv-list').innerHTML = convos.map((conv,i) => {
    const init = (conv.lead_name||'?')[0].toUpperCase();
    const msgs = (liFrom||liTo) ? conv.messages.filter(m=>_liInRange(m.date)) : conv.messages;
    const msgHtml = msgs.map(m => {
      const sent = m.automated || m.sender === conv.owner_name;
      return `<div class="msg-row ${sent?'sent':'received'}">
        <div class="msg-av">${esc((m.sender||'?')[0].toUpperCase())}</div>
        <div><div class="msg-meta">${esc(m.sender)}${m.automated?' <span class="auto-tag">[auto]</span>':''} · ${m.date}</div>
        <div class="msg-bubble">${esc(m.body)}</div></div></div>`;
    }).join('');
    return `<div class="conv-item">
      <div class="conv-header" onclick="toggleConv('li',${i})">
        <div class="av">${conv.lead_picture?`<img src="${esc(conv.lead_picture)}" onerror="this.parentElement.textContent='${init}'">`:init}</div>
        <div class="conv-info"><div class="conv-name">${esc(conv.lead_name)}</div>
          <div class="conv-occ">${esc((conv.lead_occupation||'').slice(0,60))}</div></div>
        <div class="conv-right">
          ${conv.unread>0?`<span class="unread-pill">${conv.unread}</span>`:''}
          <span style="font-size:10px;color:#64748b">${msgs.length} msgs</span>
          <span style="font-size:10px;color:#94a3b8">${esc(conv.owner_name)}</span>
          <span class="chev" id="li-chev-${i}">▼</span>
        </div>
      </div>
      <div class="conv-body" id="li-cb-${i}">${msgHtml}</div></div>`;
  }).join('') || '<div class="no-data">No conversations.</div>';
}

function renderLICamps() {
  const camps = filtLICamps();
  document.getElementById('li-camp-count').textContent = camps.length;
  document.getElementById('li-camp-body').innerHTML = camps.map(c => `<tr>
    <td style="font-weight:600;color:#1e293b">${esc(c.name)}</td>
    <td><span class="badge ${SB[c.state]||'b-init'}">${c.state}</span></td>
    <td style="color:#64748b">${c.type||'—'}</td>
    <td><b>${fmt(c.targets)}</b></td>
    <td style="color:#15803d;font-weight:700">${c.accepted_recent||0}</td>
    <td style="color:#1d4ed8;font-weight:700">${c.replies_recent||0}</td>
    <td style="color:#64748b;font-size:11px">${(c.owner_names||[]).join(', ')}</td>
    <td style="color:#94a3b8;font-size:11px">${c.created||''}</td></tr>`).join('')
    || '<tr><td colspan="8" class="no-data">No campaigns.</td></tr>';
}

function renderLILeads() {
  const leads = filtLILeads();
  document.getElementById('li-leads-count').textContent = leads.length;
  document.getElementById('li-leads-body').innerHTML = leads.map(l => `<tr>
    <td><span class="badge ${l.transition==='accepted'?'b-accepted':'b-reply'}">${l.transition.toUpperCase()}</span></td>
    <td style="font-weight:600;color:#1e293b">${esc(l.target_name)}</td>
    <td style="color:#64748b;font-size:11px">${esc((l.target_occupation||'').slice(0,50))}</td>
    <td style="font-size:11px;color:#374151">${esc(l.campaign_name)}</td>
    <td style="color:#64748b;font-size:11px">${esc(l.account_name)}</td>
    <td style="color:#94a3b8;font-size:11px">${l.timestamp}</td></tr>`).join('')
    || '<tr><td colspan="6" class="no-data">No recent leads.</td></tr>';
}

function toggleConv(ns, i) {
  const b = document.getElementById(ns+'-cb-'+i);
  const ch = document.getElementById(ns+'-chev-'+i);
  b.classList.toggle('open');
  ch.style.transform = b.classList.contains('open') ? 'rotate(180deg)' : '';
}

// LinkedIn date helpers
function setLIToday(){const t=new Date().toISOString().slice(0,10);document.getElementById('li-from').value=t;document.getElementById('li-to').value=t;renderLI()}
function setLIWeek(){const n=new Date(),d=n.getDay(),m=new Date(n);m.setDate(n.getDate()-(d===0?6:d-1));const s=new Date(m);s.setDate(m.getDate()+6);document.getElementById('li-from').value=m.toISOString().slice(0,10);document.getElementById('li-to').value=s.toISOString().slice(0,10);renderLI()}
function clearLIDate(){document.getElementById('li-from').value='';document.getElementById('li-to').value='';renderLI()}

// ── LinkedIn Segment Stat Card ────────────────────────────────────────
let _liStatTimer = null;
function loadLIStatCard(init=false) {
  clearTimeout(_liStatTimer);
  _liStatTimer = setTimeout(() => _fetchLIStatCard(init), 300);
}

async function _fetchLIStatCard(init=false) {
  const seg    = document.getElementById('li-stat-segment')?.value || 'all';
  const sub    = document.getElementById('li-stat-sub')?.value    || 'all';
  const client = smClient || 'ALL';   // sync with email client filter
  const cards  = document.getElementById('li-stat-cards');
  const hint   = document.getElementById('li-stat-loading');
  if (hint) hint.textContent = client !== 'ALL' ? `Synced with: ${client}` : 'Loading…';
  try {
    const r = await fetch(`/api/linkedin-stats?segment=${encodeURIComponent(seg)}&sub_segment=${encodeURIComponent(sub)}&client=${encodeURIComponent(client)}`);
    if (!r.ok) throw new Error('HTTP ' + r.status);
    const d = await r.json();
    if (d.error) throw new Error(d.error);

    if (init) {
      // Populate dropdowns on first load
      const segSel = document.getElementById('li-stat-segment');
      const subSel = document.getElementById('li-stat-sub');
      if (segSel && d.segment_opts) {
        segSel.innerHTML = '<option value="all">All Accounts</option>' +
          (d.segment_opts).map(o=>`<option value="${esc(String(o.id))}">${esc(o.name)}</option>`).join('');
      }
      if (subSel && d.sub_segment_opts) {
        subSel.innerHTML = '<option value="all">All Types</option>' +
          (d.sub_segment_opts).map(o=>`<option value="${esc(o.id)}">${esc(o.name)}</option>`).join('');
      }
    }

    if (cards) cards.innerHTML = `
      <div class="stat-card c-blue">
        <div class="stat-icon"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#2563eb" stroke-width="2.5" stroke-linecap="round"><path d="M17 21v-2a4 4 0 0 0-4-4H5a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M23 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/></svg></div>
        <div class="stat-val">${fmt(d.total_leads)}</div><div class="stat-lbl">Total Leads</div>
        <div class="stat-sub">Recent activity</div>
      </div>
      <div class="stat-card c-slate">
        <div class="stat-icon"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#475569" stroke-width="2.5" stroke-linecap="round"><line x1="22" y1="2" x2="11" y2="13"/><polygon points="22 2 15 22 11 13 2 9 22 2"/></svg></div>
        <div class="stat-val">${fmt(d.requests_sent)}</div><div class="stat-lbl">Requests Sent</div>
        <div class="stat-sub">Campaign targets</div>
      </div>
      <div class="stat-card c-green">
        <div class="stat-icon"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#16a34a" stroke-width="2.5" stroke-linecap="round"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg></div>
        <div class="stat-val">${d.acceptance_rate}</div><div class="stat-lbl">Acceptance Rate</div>
        <div class="stat-sub">${fmt(d.accepted)} accepted</div>
      </div>
      <div class="stat-card c-purple">
        <div class="stat-icon"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#7c3aed" stroke-width="2.5" stroke-linecap="round"><polyline points="9 17 4 12 9 7"/><path d="M20 18v-2a4 4 0 0 0-4-4H4"/></svg></div>
        <div class="stat-val">${d.reply_rate}</div><div class="stat-lbl">Reply Rate</div>
        <div class="stat-sub">${fmt(d.replies)} replies</div>
      </div>`;
    if (hint) hint.textContent = '';
  } catch(e) {
    if (cards) cards.innerHTML = `<div style="color:#dc2626;font-size:12px;padding:4px">Error: ${esc(e.message)}</div>`;
    if (hint)  hint.textContent = '';
  }
}

// LinkedIn drawer
function showLIDetail(type) {
  const cfg = {
    requests:{title:'Requests Sent — Campaigns',color:'linear-gradient(135deg,#3b82f6,#2563eb)'},
    accepted:{title:'Connections Accepted',color:'linear-gradient(135deg,#22c55e,#16a34a)'},
    messages:{title:'Messages Sent — Conversations',color:'linear-gradient(135deg,#a855f7,#7c3aed)'},
    replies: {title:'Replies Received',color:'linear-gradient(135deg,#f97316,#ea580c)'},
  };
  const c = cfg[type];
  document.getElementById('li-drawer-icon').style.background = c.color;
  document.getElementById('li-drawer-icon').innerHTML = '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#fff" stroke-width="2.5" stroke-linecap="round"><circle cx="12" cy="12" r="10"/></svg>';
  document.getElementById('li-drawer-title').textContent = c.title;
  const body = document.getElementById('li-drawer-body');
  let html = '', count = 0;
  if (type === 'requests') {
    const camps = filtLICamps(); count = camps.length;
    html = camps.length ? `<table><thead><tr><th>Campaign</th><th>State</th><th>Targets</th><th>Accepted</th><th>Replies</th><th>Owner</th></tr></thead><tbody>`
      + camps.map(c=>`<tr><td style="font-weight:600">${esc(c.name)}</td><td><span class="badge ${SB[c.state]||'b-init'}">${c.state}</span></td><td>${fmt(c.targets)}</td><td style="color:#15803d;font-weight:700">${c.accepted_recent||0}</td><td style="color:#1d4ed8;font-weight:700">${c.replies_recent||0}</td><td style="color:#64748b;font-size:11px">${(c.owner_names||[]).join(', ')}</td></tr>`).join('')+'</tbody></table>'
      : '<div style="padding:32px;text-align:center;color:#94a3b8">No campaigns.</div>';
  } else if (type === 'accepted' || type === 'replies') {
    const t = type === 'accepted' ? 'accepted' : 'reply';
    const leads = filtLILeads().filter(l=>l.transition===t); count = leads.length;
    html = leads.length ? `<table><thead><tr><th>Lead</th><th>Role</th><th>Campaign</th><th>Account</th><th>Date</th></tr></thead><tbody>`
      + leads.map(l=>`<tr><td style="font-weight:600">${esc(l.target_name)}</td><td style="color:#64748b;font-size:11px">${esc((l.target_occupation||'').slice(0,50))}</td><td style="font-size:11px">${esc(l.campaign_name)}</td><td style="color:#64748b;font-size:11px">${esc(l.account_name)}</td><td style="color:#94a3b8;font-size:11px">${l.timestamp}</td></tr>`).join('')+'</tbody></table>'
      : '<div style="padding:32px;text-align:center;color:#94a3b8">No leads found.</div>';
  } else if (type === 'messages') {
    const convos = filtLIConvos().filter(c=>c.connected||c.messages.length>0); count = convos.length;
    html = convos.length ? `<table><thead><tr><th>Lead</th><th>Role</th><th>Account</th><th>Messages</th><th>Unread</th></tr></thead><tbody>`
      + convos.map(conv=>`<tr><td style="font-weight:600">${esc(conv.lead_name)}</td><td style="color:#64748b;font-size:11px">${esc((conv.lead_occupation||'').slice(0,50))}</td><td style="color:#64748b;font-size:11px">${esc(conv.owner_name)}</td><td style="color:#6d28d9;font-weight:700">${conv.messages.length}</td><td>${conv.unread>0?`<span class="unread-pill">${conv.unread}</span>`:'—'}</td></tr>`).join('')+'</tbody></table>'
      : '<div style="padding:32px;text-align:center;color:#94a3b8">No active conversations.</div>';
  }
  body.innerHTML = html;
  document.getElementById('li-drawer-count').textContent = count;
  document.getElementById('li-drawer').classList.add('open');
  document.getElementById('li-drawer-backdrop').classList.add('open');
}
function closeLIDrawer(){ document.getElementById('li-drawer').classList.remove('open'); document.getElementById('li-drawer-backdrop').classList.remove('open'); }

// ═══════════════════════════════════════════════════════════════════
//  ZOHO
// ═══════════════════════════════════════════════════════════════════
const COMPANY_COLORS = {
  'Feaam':{'bg':'#ede9fe','color':'#6d28d9','dot':'#7c3aed'},
  'K&M Property':{'bg':'#dcfce7','color':'#15803d','dot':'#16a34a'},
  'Stanford G':{'bg':'#fff7ed','color':'#c2410c','dot':'#ea580c'},
  'Nexus':{'bg':'#dbeafe','color':'#1d4ed8','dot':'#2563eb'},
  'Kendra':{'bg':'#fce7f3','color':'#9d174d','dot':'#be185d'},
  'Unknown':{'bg':'#f1f5f9','color':'#475569','dot':'#94a3b8'},
};

async function loadZoho() {
  try {
    const r = await fetch('/api/zoho-stats');
    _zoho = await r.json();
    renderZoho();
  } catch(e) {
    document.getElementById('zoho-body').innerHTML = '<span style="color:#ef4444;font-size:12px">Failed to load sync stats</span>';
  }
}

function renderZoho() {
  if (!_zoho) return;
  document.getElementById('zoho-ts').textContent = _zoho.last_sync ? 'Last sync: '+_zoho.last_sync : 'Never synced';
  const mkChip = (lbl, val, bg, color) =>
    `<div class="zoho-chip" style="background:${bg}"><div class="zoho-chip-val" style="color:${color}">${val}</div><div class="zoho-chip-lbl" style="color:${color}">${lbl}</div></div>`;
  const companies = _zoho.by_company || {};
  const compHTML = Object.entries(companies).map(([name, count]) => {
    const col = COMPANY_COLORS[name] || COMPANY_COLORS['Unknown'];
    return `<div class="zoho-chip" style="background:${col.bg}">
      <div style="display:flex;align-items:center;gap:5px;margin-bottom:2px"><div style="width:8px;height:8px;border-radius:50%;background:${col.dot}"></div><span style="font-size:10px;font-weight:700;color:${col.color}">${esc(name)}</span></div>
      <div class="zoho-chip-val" style="color:${col.color}">${count}</div></div>`;
  }).join('');
  document.getElementById('zoho-body').innerHTML =
    mkChip('Total Synced', _zoho.total||0, 'linear-gradient(135deg,#e0f2fe,#bae6fd)', '#0369a1') +
    mkChip('Created', _zoho.created||0, 'linear-gradient(135deg,#dcfce7,#bbf7d0)', '#15803d') +
    mkChip('Updated', _zoho.updated||0, 'linear-gradient(135deg,#fef3c7,#fde68a)', '#92400e') +
    compHTML;
}

async function triggerZohoSync() {
  const btn = document.getElementById('zoho-sync-btn');
  btn.disabled = true; btn.textContent = 'Syncing…';
  try {
    const r = await fetch('/api/zoho-sync', {method:'POST'});
    const d = await r.json();
    if (d.ok) await loadZoho();
    else alert('Sync failed: '+(d.error||'unknown'));
  } catch(e) { alert('Sync error: '+e.message); }
  finally { btn.disabled=false; btn.innerHTML='<svg width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg> Sync Now'; }
}

// ═══════════════════════════════════════════════════════════════════
//  EMAIL DATA
// ═══════════════════════════════════════════════════════════════════
let _emailPollTimer = null;

async function loadEmail() {
  const hint = document.getElementById('sm-loading-hint');
  const err  = document.getElementById('sm-err');
  err.style.display = 'none';

  // Show a progress banner
  let banner = document.getElementById('email-fetch-banner');
  if (!banner) {
    banner = document.createElement('div');
    banner.id = 'email-fetch-banner';
    banner.style.cssText = 'background:#fef9c3;border:1px solid #fde047;border-radius:8px;padding:10px 16px;margin-bottom:12px;font-size:13px;color:#78350f;display:flex;align-items:center;gap:10px;';
    const smSection = document.getElementById('sm-loading-hint')?.closest('.card') || document.getElementById('sm-loading-hint')?.parentElement;
    if (smSection) smSection.prepend(banner);
  }

  hint.textContent = '';
  try {
    const r = await fetch('/api/email');
    if (!r.ok) throw new Error('HTTP '+r.status);
    SM = await r.json();
    if (SM.error) throw new Error(SM.error);

    if (SM.loading) {
      // Show sample data immediately + progress
      const p = SM.progress || {};
      const pct = p.total ? Math.round(p.done/p.total*100) : 0;
      banner.innerHTML = `<span>⏳</span><span><b>Fetching live data…</b> ${p.phase||''} ${p.total ? `(${p.done}/${p.total} — ${pct}%)` : ''} — <em>Showing sample preview below</em></span>`;
      banner.style.display = 'flex';
      populateSMClients();
      resetEmailFilters();
      // Poll status every 8 seconds until done
      if (_emailPollTimer) clearTimeout(_emailPollTimer);
      _emailPollTimer = setTimeout(() => pollEmailUntilDone(), 8000);
    } else {
      banner.style.display = 'none';
      if (_emailPollTimer) { clearTimeout(_emailPollTimer); _emailPollTimer = null; }
      document.getElementById('sm-ts').textContent = 'Email: '+new Date().toLocaleTimeString();
      document.getElementById('sm-cached').style.display = SM.cached ? 'inline-block' : 'none';
      populateSMClients();
      resetEmailFilters();
    }
  } catch(e) {
    banner.style.display = 'none';
    err.textContent = 'Email error: '+e.message;
    err.style.display = 'block';
  }
}

async function pollEmailUntilDone() {
  try {
    const r = await fetch('/api/email/status');
    const s = await r.json();
    const banner = document.getElementById('email-fetch-banner');
    if (s.loading) {
      const p = s.progress || {};
      const pct = p.total ? Math.round(p.done/p.total*100) : 0;
      if (banner) banner.innerHTML = `<span>⏳</span><span><b>Fetching live data…</b> ${p.phase||''} ${p.total ? `(${p.done}/${p.total} — ${pct}%)` : ''} — <em>Showing sample preview below</em></span>`;
      _emailPollTimer = setTimeout(() => pollEmailUntilDone(), 8000);
    } else {
      // Data is ready — reload with real data
      if (banner) banner.innerHTML = '<span>✅</span><span><b>Live data ready!</b> Loading now…</span>';
      setTimeout(() => loadEmail(), 500);
    }
  } catch(e) {
    _emailPollTimer = setTimeout(() => pollEmailUntilDone(), 8000);
  }
}

function populateSMClients() {
  const clients = [...new Set([
    ...SM.parent_analytics.map(r=>r.client),
    ...SM.sub_analytics.map(r=>r.parent)
  ])].sort();
  const sel = document.getElementById('sm-client');
  sel.innerHTML = '<option value="ALL">All Clients</option>' + clients.map(c=>`<option value="${esc(c)}">${esc(c)}</option>`).join('');
}

function smRecompute(leads) {
  const total=leads.length, opened=leads.filter(l=>l.open_time).length,
        clicked=leads.filter(l=>l.click_time).length, replied=leads.filter(l=>l.reply_time).length,
        bounced=leads.filter(l=>l.is_bounced).length, unsubscribed=leads.filter(l=>l.is_unsubscribed).length,
        positive=leads.filter(l=>POSITIVE_CATS.has((l.category||'').toLowerCase())).length;
  const r=n=>total?+(n/total*100).toFixed(2):0;
  return {total,opened,clicked,replied,bounced,unsubscribed,positive,
          open_rate:r(opened),click_rate:r(clicked),reply_rate:r(replied),positive_rate:r(positive)};
}

function applyEmailFilters() {
  if (!SM) return;
  smClient = document.getElementById('sm-client').value;
  smFrom   = document.getElementById('sm-from').value;
  smTo     = document.getElementById('sm-to').value;

  // Base client + date filter
  smParentFilt = SM.parent_analytics.filter(r => smClient==='ALL' || r.client===smClient);
  smSubFilt = SM.sub_analytics.map(row => {
    if (smClient!=='ALL' && row.parent!==smClient) return null;
    let leads = row.leads || [];
    if (smFrom||smTo) leads = leads.filter(l => {
      const d=(l.sent_time||l.open_time||'').substring(0,10);
      if (!d) return false;
      if (smFrom&&d<smFrom) return false;
      if (smTo&&d>smTo) return false;
      return true;
    });
    return {...row, leads, ...smRecompute(leads)};
  }).filter(r=>r!==null);

  // Drill-down filters (sub-sub dashboard)
  if (drillClient) {
    smParentFilt = smParentFilt.filter(r => r.client === drillClient);
    smSubFilt    = smSubFilt.filter(r => r.parent === drillClient);
  }
  if (drillParent) {
    smParentFilt = smParentFilt.filter(r => r.raw_name === drillParent);
    smSubFilt    = smSubFilt.filter(r => r.parent === drillClient);
  }
  if (drillSub) {
    smSubFilt = smSubFilt.filter(r => r.subsequence === drillSub);
  }

  const parts=[];
  if (smClient!=='ALL') parts.push(smClient);
  if (drillClient && drillClient!==smClient) parts.push(drillClient);
  if (drillParent) parts.push(drillParent);
  if (drillSub)    parts.push(drillSub);
  if (smFrom||smTo) parts.push((smFrom||'…')+' → '+(smTo||'…'));
  document.getElementById('sm-filter-status').textContent = parts.length?'Filtered: '+parts.join(' › '):'';
  renderEmail();
  loadLIStatCard();  // keep LinkedIn stat card in sync with email client filter
}

function resetEmailFilters() {
  document.getElementById('sm-client').value='ALL';
  document.getElementById('sm-from').value='';
  document.getElementById('sm-to').value='';
  document.getElementById('sm-filter-status').textContent='';
  smParentFilt = (SM?.parent_analytics||[]).slice();
  smSubFilt = (SM?.sub_analytics||[]).map(r=>({...r,...smRecompute(r.leads||[])}));
  renderEmail();
}

function renderEmail() {
  renderSMCombinedCards(); renderSMPipelineChart(); renderSMCombinedTable();
  renderSMMainCards(); renderSMOpenedChart(); renderSMSubByClientChart();
  renderSMMainTable(); renderSMSubCards(); renderSMSubCharts(); renderSMSubTable();
  renderPositiveLeads();
  document.getElementById('sm-camp-count').textContent = smParentFilt.length;
  document.getElementById('sm-sub-count').textContent  = smSubFilt.length;
}

// ── Combined Overview ──
function renderSMCombinedCards() {
  const totSent   = smParentFilt.reduce((s,r)=>s+r.total,0);
  const totPOpen  = smParentFilt.reduce((s,r)=>s+r.opened,0);
  const totPReply = smParentFilt.reduce((s,r)=>s+r.replied,0);
  const totSub    = smParentFilt.reduce((s,r)=>s+r.added_to_sub,0);
  const totPos    = smParentFilt.reduce((s,r)=>s+r.positive,0);
  const totSOpen  = smSubFilt.reduce((s,r)=>s+r.opened,0);
  const totSClick = smSubFilt.reduce((s,r)=>s+r.clicked,0);
  const totSReply = smSubFilt.reduce((s,r)=>s+r.replied,0);
  const totSLeads = smSubFilt.reduce((s,r)=>s+r.total,0);
  document.getElementById('sm-combined-cards').innerHTML = `
    <div class="stat-card c-slate"><div class="stat-lbl">Total Sent</div><div class="stat-val">${fmt(totSent)}</div><div class="stat-sub">${smParentFilt.length} campaigns</div></div>
    <div class="stat-card c-green"><div class="stat-lbl">Lead. Opened</div><div class="stat-val">${fmt(totPOpen)}</div><div class="stat-sub">${pct(totPOpen,totSent)} of sent</div></div>
    <div class="stat-card c-purple"><div class="stat-lbl">Lead. Replied</div><div class="stat-val">${fmt(totPReply)}</div><div class="stat-sub">${pct(totPReply,totSent)} of sent</div></div>
    <div class="stat-card c-blue"><div class="stat-lbl">Added to Sub</div><div class="stat-val">${fmt(totSub)}</div><div class="stat-sub">${pct(totSub,totSent)} of sent</div></div>
    <div class="stat-card c-emerald"><div class="stat-lbl">Sub Opened</div><div class="stat-val">${fmt(totSOpen)}</div><div class="stat-sub">${pct(totSOpen,totSLeads)} of sub leads</div></div>
    <div class="stat-card c-cyan"><div class="stat-lbl">Sub Clicked</div><div class="stat-val">${fmt(totSClick)}</div><div class="stat-sub">${pct(totSClick,totSLeads)} of sub leads</div></div>
    <div class="stat-card c-amber"><div class="stat-lbl">Sub Replied</div><div class="stat-val">${fmt(totSReply)}</div><div class="stat-sub">${pct(totSReply,totSLeads)} of sub leads</div></div>
    <div class="stat-card c-rose"><div class="stat-lbl">Total Positive</div><div class="stat-val">${fmt(totPos)}</div><div class="stat-sub">${pct(totPos,totSent)} conversion</div></div>`;
}

function renderSMPipelineChart() {
  const cm={};
  smParentFilt.forEach(r=>{if(!cm[r.client])cm[r.client]={sent:0,pOpen:0,pReply:0,sub:0,sOpen:0,sReply:0};cm[r.client].sent+=r.total;cm[r.client].pOpen+=r.opened;cm[r.client].pReply+=r.replied;cm[r.client].sub+=r.added_to_sub});
  smSubFilt.forEach(r=>{if(!cm[r.parent])cm[r.parent]={sent:0,pOpen:0,pReply:0,sub:0,sOpen:0,sReply:0};cm[r.parent].sOpen+=r.opened;cm[r.parent].sReply+=r.replied});
  const labels=Object.keys(cm), vals=Object.values(cm);
  if (charts['smPipeline']) charts['smPipeline'].destroy();
  charts['smPipeline'] = new Chart(document.getElementById('sm-pipeline-chart'),{
    type:'bar', data:{labels, datasets:[
      {label:'Sent',data:vals.map(d=>d.sent),backgroundColor:'#94a3b8',borderRadius:3},
      {label:'Lead. Opened',data:vals.map(d=>d.pOpen),backgroundColor:'#22d3ee',borderRadius:3},
      {label:'Lead. Replied',data:vals.map(d=>d.pReply),backgroundColor:'#c084fc',borderRadius:3},
      {label:'Added to Sub',data:vals.map(d=>d.sub),backgroundColor:'#38bdf8',borderRadius:3},
      {label:'Sub Opened',data:vals.map(d=>d.sOpen),backgroundColor:'#4ade80',borderRadius:3},
      {label:'Sub Replied',data:vals.map(d=>d.sReply),backgroundColor:'#fb923c',borderRadius:3},
    ]},
    options:{responsive:true,maintainAspectRatio:false,plugins:{legend:{labels:{color:'#64748b',font:{size:10}}}},scales:{x:{ticks:{color:'#64748b',font:{size:10}},grid:{color:'#f8fafc'}},y:{ticks:{color:'#64748b'},grid:{color:'#f1f5f9'}}}}
  });
}

function renderSMCombinedTable() {
  const cm={};
  smParentFilt.forEach(r=>{if(!cm[r.client])cm[r.client]={sent:0,pOpen:0,pReply:0,pos:0,sub:0,sLeads:0,sOpen:0,sClick:0,sReply:0};cm[r.client].sent+=r.total;cm[r.client].pOpen+=r.opened;cm[r.client].pReply+=r.replied;cm[r.client].pos+=r.positive;cm[r.client].sub+=r.added_to_sub});
  smSubFilt.forEach(r=>{if(!cm[r.parent])cm[r.parent]={sent:0,pOpen:0,pReply:0,pos:0,sub:0,sLeads:0,sOpen:0,sClick:0,sReply:0};cm[r.parent].sLeads+=r.total;cm[r.parent].sOpen+=r.opened;cm[r.parent].sClick+=r.clicked;cm[r.parent].sReply+=r.replied});
  const rows=Object.entries(cm);
  window._smCombinedRows=rows;
  if(!rows.length){document.getElementById('sm-combined-table').innerHTML='<div class="no-data">No data.</div>';return}
  let h=`<table><thead><tr><th>Client</th><th>Sent</th><th>Lead. Opened</th><th>Lead. Replied</th><th>Added to Sub</th><th>Sub Leads</th><th>Sub Opened</th><th>Sub Clicked</th><th>Sub Replied</th><th>Positive</th></tr></thead><tbody>`;
  rows.forEach(([client,d],idx)=>{
    const mk=(val,tot,type,col)=>{const inner=`${val} <span class="rate">(${pct(val,tot)})</span>`;return val>0?`<span class="clickable"${col?' style="color:'+col+'"':''} ondblclick="openSMCombinedModal(${idx},'${type}')">${inner}</span>`:inner};
    h+=`<tr><td><b>${esc(client)}</b></td><td>${d.sent}</td><td>${mk(d.pOpen,d.sent,'pOpen','')}</td><td>${mk(d.pReply,d.sent,'pReply','')}</td><td>${mk(d.sub,d.sent,'sub','#2563eb')}</td><td>${d.sLeads}</td><td>${mk(d.sOpen,d.sLeads,'sOpen','')}</td><td>${mk(d.sClick,d.sLeads,'sClick','')}</td><td>${mk(d.sReply,d.sLeads,'sReply','')}</td><td>${mk(d.pos,d.sent,'pos','#059669')}</td></tr>`;
  });
  h+='</tbody></table>';
  document.getElementById('sm-combined-table').innerHTML=h;
}

// ── Main Campaign Analytics ──
function renderSMMainCards() {
  const tot=smParentFilt.reduce((s,r)=>s+r.total,0);
  const opn=smParentFilt.reduce((s,r)=>s+r.opened,0);
  const rep=smParentFilt.reduce((s,r)=>s+r.replied,0);
  const sub=smParentFilt.reduce((s,r)=>s+r.added_to_sub,0);
  document.getElementById('sm-main-cards').innerHTML=`
    <div class="stat-card c-slate"><div class="stat-lbl">Total Sent</div><div class="stat-val">${fmt(tot)}</div><div class="stat-sub">${smParentFilt.length} campaigns</div></div>
    <div class="stat-card c-green"><div class="stat-lbl">Lead Opened</div><div class="stat-val">${fmt(opn)}</div><div class="stat-sub">${pct(opn,tot)} of sent</div></div>
    <div class="stat-card c-purple"><div class="stat-lbl">Lead Replied</div><div class="stat-val">${fmt(rep)}</div><div class="stat-sub">${pct(rep,tot)} of sent</div></div>
    <div class="stat-card c-blue"><div class="stat-lbl">Added to Sub</div><div class="stat-val">${fmt(sub)}</div><div class="stat-sub">${pct(sub,tot)} of sent</div></div>`;
}

function mkChart(id, type, labels, datasets, opts={}) {
  if (charts[id]) charts[id].destroy();
  charts[id] = new Chart(document.getElementById(id), {
    type, data:{labels, datasets},
    options:{responsive:true,maintainAspectRatio:false,
             plugins:{legend:{display:false},...(opts.plugins||{})},
             scales:{x:{ticks:{color:'#64748b',font:{size:10}},grid:{color:'#f8fafc'}},
                     y:{ticks:{color:'#64748b'},grid:{color:'#f1f5f9'}},...(opts.scales||{})},
             ...opts}});
}

function renderSMOpenedChart() {
  const labels=smParentFilt.map(r=>r.client), data=smParentFilt.map(r=>r.opened);
  mkChart('sm-opened-chart','doughnut',labels,
    [{data,backgroundColor:['#3b82f6','#22c55e','#a855f7','#f59e0b','#06b6d4','#ec4899','#14b8a6','#f97316'],borderWidth:2}],
    {plugins:{legend:{display:true,position:'right',labels:{color:'#64748b',font:{size:10}}}},scales:{x:{display:false},y:{display:false}}});
}

function renderSMSubByClientChart() {
  const labels=smParentFilt.map(r=>r.client), data=smParentFilt.map(r=>r.added_to_sub);
  mkChart('sm-sub-chart','doughnut',labels,
    [{data,backgroundColor:['#38bdf8','#4ade80','#fb923c','#c084fc','#fbbf24','#f87171','#34d399','#818cf8'],borderWidth:2}],
    {plugins:{legend:{display:true,position:'right',labels:{color:'#64748b',font:{size:10}}}},scales:{x:{display:false},y:{display:false}}});
}

function renderSMMainTable() {
  window._smParentData = smParentFilt;
  if (!pagers['sm-main-table']) mkPager('sm-main-table', 5);
  const hdr = `<tr><th>Client</th><th>Campaign</th><th>Status</th><th>Total Sent</th><th>Opened</th><th>Replied</th><th>Added to Sub</th><th>Positive</th></tr>`;
  pagers['sm-main-table'].init(smParentFilt, hdr, (r, i) => {
    const mk = (val,tot,type) => val>0
      ? `<span class="clickable" ondblclick="openSMParentModal(${i},'${type}')">${val} <span class="rate">(${pct(val,tot)})</span></span>`
      : val;
    return `<tr>
      <td><b class="bc-item" style="cursor:pointer" onclick="setDrill('${r.client}',null,null)">${esc(r.client)}</b></td>
      <td style="color:#64748b;font-size:11px;cursor:pointer" class="bc-item" onclick="setDrill('${r.client}','${r.raw_name}',null)">${esc(r.raw_name)}</td>
      <td><span class="badge b-init">${esc(r.status)}</span></td>
      <td>${r.total}</td>
      <td>${mk(r.opened,r.total,'opened')}</td>
      <td>${mk(r.replied,r.total,'replied')}</td>
      <td style="color:#2563eb;font-weight:600">${mk(r.added_to_sub,r.total,'sub')}</td>
      <td style="color:#059669;font-weight:600">${mk(r.positive,r.total,'positive')}</td>
    </tr>`;
  });
}

// ── Subsequence Analytics ──
function renderSMSubCards() {
  const tot=smSubFilt.reduce((s,r)=>s+r.total,0);
  const opn=smSubFilt.reduce((s,r)=>s+r.opened,0);
  const clk=smSubFilt.reduce((s,r)=>s+r.clicked,0);
  const rep=smSubFilt.reduce((s,r)=>s+r.replied,0);
  document.getElementById('sm-sub-cards').innerHTML=`
    <div class="stat-card c-blue"><div class="stat-lbl">Total Sub Leads</div><div class="stat-val">${fmt(tot)}</div><div class="stat-sub">${smSubFilt.length} subsequences</div></div>
    <div class="stat-card c-emerald"><div class="stat-lbl">Sub Opened</div><div class="stat-val">${fmt(opn)}</div><div class="stat-sub">${pct(opn,tot)} of sub leads</div></div>
    <div class="stat-card c-cyan"><div class="stat-lbl">Sub Clicked</div><div class="stat-val">${fmt(clk)}</div><div class="stat-sub">${pct(clk,tot)} of sub leads</div></div>
    <div class="stat-card c-amber"><div class="stat-lbl">Sub Replied</div><div class="stat-val">${fmt(rep)}</div><div class="stat-sub">${pct(rep,tot)} of sub leads</div></div>`;
}

function renderSMSubCharts() {
  const MAX=15;
  const sorted=[...smSubFilt].sort((a,b)=>b.total-a.total).slice(0,MAX);
  const byOpen=[...smSubFilt].sort((a,b)=>b.open_rate-a.open_rate).slice(0,MAX);
  const byClick=[...smSubFilt].sort((a,b)=>b.click_rate-a.click_rate).slice(0,MAX);
  const byReply=[...smSubFilt].sort((a,b)=>b.reply_rate-a.reply_rate).slice(0,MAX);
  const bar=(id,data,key,col)=>mkChart(id,'bar',data.map(r=>r.subsequence.slice(0,20)),
    [{data:data.map(r=>r[key]),backgroundColor:col,borderRadius:3}],
    {indexAxis:'y',scales:{x:{ticks:{color:'#64748b'},grid:{color:'#f1f5f9'}},y:{ticks:{color:'#64748b',font:{size:9}}}}});
  bar('sm-open-rate-chart',byOpen,'open_rate','#22c55e');
  bar('sm-click-rate-chart',byClick,'click_rate','#3b82f6');
  bar('sm-reply-rate-chart',byReply,'reply_rate','#a855f7');
  mkChart('sm-leads-chart','bar',sorted.map(r=>r.subsequence.slice(0,20)),
    [{data:sorted.map(r=>r.total),backgroundColor:'#38bdf8',borderRadius:3}],
    {scales:{x:{ticks:{color:'#64748b',font:{size:9}},grid:{color:'#f8fafc'}},y:{ticks:{color:'#64748b'},grid:{color:'#f1f5f9'}}}});
}

function renderSMSubTable() {
  window._smSubData = smSubFilt;
  if (!pagers['sm-sub-table']) mkPager('sm-sub-table', 5);
  const hdr = `<tr><th>Client</th><th>Subsequence</th><th>Status</th><th>Total</th><th>Opened</th><th>Clicked</th><th>Replied</th><th>Bounced</th></tr>`;
  pagers['sm-sub-table'].init(smSubFilt, hdr, (r, i) => {
    const mk = (val,tot,type) => val>0
      ? `<span class="clickable" ondblclick="openSMSubModal(${i},'${type}')">${val} <span class="rate">(${pct(val,tot)})</span></span>`
      : val;
    return `<tr>
      <td><b class="bc-item" style="cursor:pointer" onclick="setDrill('${r.parent}',null,null)">${esc(r.parent)}</b></td>
      <td style="color:#374151;cursor:pointer" class="bc-item" onclick="setDrill('${r.parent}',null,'${r.subsequence}')">${esc(r.subsequence)}</td>
      <td><span class="badge b-init">${esc(r.status)}</span></td>
      <td>${r.total}</td>
      <td>${mk(r.opened,r.total,'opened')}</td>
      <td>${mk(r.clicked,r.total,'clicked')}</td>
      <td>${mk(r.replied,r.total,'replied')}</td>
      <td>${r.bounced||0}</td>
    </tr>`;
  });
}

// ── Email Modals ──
function openSMCombinedModal(rowIdx,type){
  const [client]=window._smCombinedRows[rowIdx];
  let leads=[];
  if(type==='pOpen')      smParentFilt.filter(r=>r.client===client).forEach(r=>(r.leads||[]).forEach(l=>{if(l.open_time)leads.push(l)}));
  else if(type==='pReply')smParentFilt.filter(r=>r.client===client).forEach(r=>(r.leads||[]).forEach(l=>{if(l.reply_time)leads.push(l)}));
  else if(type==='pos')   smParentFilt.filter(r=>r.client===client).forEach(r=>(r.leads||[]).forEach(l=>{if(POSITIVE_CATS.has((l.category||'').toLowerCase()))leads.push(l)}));
  else if(type==='sub'){const seen=new Set();smSubFilt.filter(r=>r.parent===client).forEach(r=>(r.leads||[]).forEach(l=>{if(!seen.has(l.email)){seen.add(l.email);leads.push(l)}}))}
  else if(type==='sOpen') smSubFilt.filter(r=>r.parent===client).forEach(r=>(r.leads||[]).forEach(l=>{if(l.open_time)leads.push(l)}));
  else if(type==='sClick')smSubFilt.filter(r=>r.parent===client).forEach(r=>(r.leads||[]).forEach(l=>{if(l.click_time)leads.push(l)}));
  else if(type==='sReply')smSubFilt.filter(r=>r.parent===client).forEach(r=>(r.leads||[]).forEach(l=>{if(l.reply_time)leads.push(l)}));
  const labels={pOpen:'Lead. Opened',pReply:'Lead. Replied',pos:'Positive',sub:'Added to Sub',sOpen:'Sub Opened',sClick:'Sub Clicked',sReply:'Sub Replied'};
  openSMModal((labels[type]||type)+' — '+client, leads);
}
function openSMParentModal(idx,type){
  const r=window._smParentData[idx];
  const labels={opened:'Opened',replied:'Replied',sub:'Added to Sub',positive:'Positive'};
  let leads=(r.leads||[]).filter(l=>
    type==='opened'?l.open_time:type==='replied'?l.reply_time:
    type==='sub'?l.open_time:POSITIVE_CATS.has((l.category||'').toLowerCase()));
  openSMModal((labels[type]||type)+' — '+r.client, leads);
}
function openSMSubModal(idx,type){
  const r=window._smSubData[idx];
  const labels={opened:'Opened',clicked:'Clicked',replied:'Replied'};
  let leads=(r.leads||[]).filter(l=>type==='opened'?l.open_time:type==='clicked'?l.click_time:l.reply_time);
  openSMModal((labels[type]||type)+' — '+r.subsequence, leads);
}

function openSMModal(title, leads) {
  document.getElementById('sm-modal-title').textContent = title;
  window._smModalLeads = leads;
  const listEl = document.getElementById('sm-lead-list');
  if (!leads.length) {
    listEl.innerHTML='<div style="padding:20px;text-align:center;color:#94a3b8;font-size:12px">No leads found.</div>';
    document.getElementById('sm-msg-panel').innerHTML='<div class="placeholder">No leads found.</div>';
  } else {
    listEl.innerHTML = leads.map((l,i)=>`<div class="lead-item" id="sm-li-${i}" onclick="selectSMLead(${i})">
      <div class="lead-name">${esc(l.name||l.email)}</div>
      <div class="lead-email">${esc(l.email)}</div>
      <div class="lead-meta">${l.category?`<span style="background:#dcfce7;color:#16a34a;padding:1px 5px;border-radius:3px;font-size:9px;font-weight:700">${esc(l.category)}</span>`:''}</div>
    </div>`).join('');
    selectSMLead(0);
  }
  document.getElementById('sm-modal').classList.add('open');
}

function selectSMLead(i) {
  document.querySelectorAll('.lead-item').forEach(el=>el.classList.remove('active'));
  const el=document.getElementById('sm-li-'+i);
  if(el) el.classList.add('active');
  const l=window._smModalLeads[i];
  if(!l){document.getElementById('sm-msg-panel').innerHTML='<div class="placeholder">No messages.</div>';return}
  const msgs=(l.messages||[]);
  const panel=document.getElementById('sm-msg-panel');
  if(!msgs.length){panel.innerHTML='<div class="placeholder">No message history for this lead.</div>';return}
  panel.innerHTML='<div>'+msgs.map(m=>{
    const isReply=m.is_reply;
    return `<div class="sm-msg-bubble${isReply?' inbound':''}">
      <div class="sm-msg-meta">
        <div><span class="${isReply?'dir-in':'dir-out'}">${isReply?'REPLY':'SENT'}</span>
          <span class="sm-msg-seq">Seq ${m.seq||0}</span></div>
        <div class="sm-msg-date">${(m.sent_time||'').substring(0,10)}</div>
      </div>
      ${m.subject?`<div class="sm-msg-subj">${esc(m.subject)}</div>`:''}
      <div class="sm-msg-body">${esc(m.body||'')}</div>
      <div class="sm-tags">
        ${m.open_time?'<span class="sm-tag tag-open">Opened</span>':''}
        ${m.click_time?'<span class="sm-tag tag-click">Clicked</span>':''}
        ${m.is_reply?'<span class="sm-tag tag-reply">Replied</span>':''}
      </div>
    </div>`;
  }).join('')+'</div>';
}

function closeSMModal(){ document.getElementById('sm-modal').classList.remove('open') }
document.getElementById('sm-modal').addEventListener('click', e=>{ if(e.target===e.currentTarget) closeSMModal() });
document.addEventListener('keydown', e=>{ if(e.key==='Escape'){closeSMModal();closeLIDrawer();closeShareModal()} });

// ═══════════════════════════════════════════════════════════════════
//  POSITIVE LEADS TABLE
// ═══════════════════════════════════════════════════════════════════
const _POS_CATS = new Set(["interested","meeting booked","positive","meeting request","will buy","warm","demo request"]);

function renderPositiveLeads() {
  const leads = [];
  for (const campaign of smParentFilt) {
    for (const l of (campaign.leads || [])) {
      if (_POS_CATS.has((l.category||'').toLowerCase())) {
        const dates = [l.reply_time, l.open_time, l.click_time, l.sent_time].filter(Boolean).sort().reverse();
        leads.push({
          name:     l.name || l.email || 'Unknown',
          company:  campaign.client || '',
          source:   campaign.raw_name || '',
          summary:  ((l.messages||[]).slice(-1)[0]?.body||'—').substring(0, 120),
          status:   l.category || '',
          last_act: (dates[0]||'').substring(0,10) || '—',
        });
      }
    }
  }
  const countEl = document.getElementById('pos-count');
  const tblCount = document.getElementById('pos-table-count');
  if (countEl)  countEl.textContent  = leads.length;
  if (tblCount) tblCount.textContent = leads.length;

  if (!pagers['pos-leads-table']) mkPager('pos-leads-table', 5);
  const hdr = `<tr><th>Name</th><th>Company</th><th>Source Campaign</th><th>Conversation Summary</th><th>Status</th><th>Last Activity</th></tr>`;
  pagers['pos-leads-table'].init(leads, hdr, (l) => `<tr>
    <td style="font-weight:600;color:#1e293b">${esc(l.name)}</td>
    <td><b>${esc(l.company)}</b></td>
    <td style="font-size:11px;color:#64748b">${esc(l.source)}</td>
    <td style="font-size:11px;color:#374151;max-width:220px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${esc(l.summary)}">${esc(l.summary)}</td>
    <td><span class="badge b-accepted">${esc(l.status)}</span></td>
    <td style="color:#94a3b8;font-size:11px">${esc(l.last_act)}</td>
  </tr>`);
}

// ═══════════════════════════════════════════════════════════════════
//  EXPORT & SHARE
// ═══════════════════════════════════════════════════════════════════
async function downloadPDF() {
  const btn = document.getElementById('pdf-btn');
  btn.disabled = true; btn.textContent = '…';
  try {
    const body = SM ? {parent_analytics: smParentFilt, sub_analytics: smSubFilt} : {};
    const r = await fetch('/api/export/pdf', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify(body),
    });
    if (!r.ok) { const d=await r.json(); throw new Error(d.error||'Export failed'); }
    const blob = await r.blob();
    _triggerDownload(blob, `analytics_${_today()}.pdf`);
  } catch(e) { alert('PDF export failed: ' + e.message); }
  finally { btn.disabled=false; btn.innerHTML='&#x21E9; PDF'; }
}

async function downloadDOCX() {
  const btn = document.getElementById('docx-btn');
  btn.disabled = true; btn.textContent = '…';
  try {
    const body = SM ? {parent_analytics: smParentFilt, sub_analytics: smSubFilt} : {};
    const r = await fetch('/api/export/docx', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify(body),
    });
    if (!r.ok) { const d=await r.json(); throw new Error(d.error||'Export failed'); }
    const blob = await r.blob();
    _triggerDownload(blob, `analytics_${_today()}.docx`);
  } catch(e) { alert('DOCX export failed: ' + e.message); }
  finally { btn.disabled=false; btn.innerHTML='&#x21E9; DOCX'; }
}

function _triggerDownload(blob, filename) {
  const url = URL.createObjectURL(blob);
  const a   = document.createElement('a');
  a.href = url; a.download = filename; a.click();
  setTimeout(() => URL.revokeObjectURL(url), 2000);
}

function _today() { return new Date().toISOString().slice(0,10).replace(/-/g,''); }

function openShareModal() {
  document.getElementById('share-email-inp').value = '';
  const msg = document.getElementById('share-msg');
  msg.textContent = ''; msg.className = 'share-msg';
  document.getElementById('share-overlay').classList.add('open');
}

function closeShareModal() {
  document.getElementById('share-overlay').classList.remove('open');
}

document.getElementById('share-overlay').addEventListener('click', e => {
  if (e.target === e.currentTarget) closeShareModal();
});

async function sendShareEmail() {
  const email  = document.getElementById('share-email-inp').value.trim();
  const msgEl  = document.getElementById('share-msg');
  const sendBtn = document.getElementById('share-send-btn');

  if (!email) {
    msgEl.textContent = 'Please enter an email address.';
    msgEl.className = 'share-msg share-err'; return;
  }

  sendBtn.disabled = true; sendBtn.textContent = 'Sending…';
  msgEl.textContent = 'Generating PDF and sending…';
  msgEl.className = 'share-msg';
  msgEl.style.display = 'block';

  try {
    const r = await fetch('/api/share/email', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({email}),
    });
    const d = await r.json();
    if (d.ok) {
      msgEl.textContent = `Report sent to ${email} successfully!`;
      msgEl.className = 'share-msg share-ok';
    } else {
      throw new Error(d.error || 'Unknown error');
    }
  } catch(e) {
    msgEl.textContent = 'Failed: ' + e.message;
    msgEl.className = 'share-msg share-err';
  } finally {
    sendBtn.disabled = false; sendBtn.textContent = 'Send Report';
  }
}
</script>
</body>
</html>"""

# ── Per-client compact sub-dashboard ──────────────────────────────────

CLIENT_DASHBOARD_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{{client.name}} — Client Dashboard</title>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:#f0f2f5;color:#1e293b;min-height:100vh}
a{color:inherit}

/* Topbar */
.topbar{background:linear-gradient(135deg,#1e1b4b 0%,#3730a3 55%,#4f46e5 100%);padding:0 28px;display:flex;align-items:center;justify-content:space-between;height:56px;position:sticky;top:0;z-index:200;box-shadow:0 4px 20px rgba(67,56,202,.4)}
.logo{font-size:16px;font-weight:800;color:#fff;display:flex;align-items:center;gap:10px}
.logo-mark{width:28px;height:28px;border-radius:8px;background:linear-gradient(135deg,#818cf8,#a78bfa);display:flex;align-items:center;justify-content:center;color:#fff;font-weight:900}
.refresh-btn{background:rgba(255,255,255,.12);color:#fff;border:1px solid rgba(255,255,255,.22);padding:7px 16px;border-radius:8px;font-size:12px;font-weight:600;cursor:pointer}
.refresh-btn:hover{background:rgba(255,255,255,.22)}
.act-btn{padding:7px 14px;border-radius:8px;font-size:12px;font-weight:700;cursor:pointer;border:none;color:#fff;display:inline-flex;align-items:center;gap:5px;font-family:inherit}
.act-pdf{background:#dc2626} .act-pdf:hover{background:#b91c1c}
.act-docx{background:#2563eb} .act-docx:hover{background:#1d4ed8}
.act-share{background:#16a34a} .act-share:hover{background:#15803d}
.act-btn:disabled{opacity:.5;cursor:not-allowed}

/* Tab nav */
.client-tabs{background:#fff;border-bottom:1px solid #e2e8f0;padding:10px 28px;display:flex;gap:6px;flex-wrap:wrap;align-items:center;position:sticky;top:56px;z-index:150}
.client-tabs .tlbl{font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:1px;margin-right:6px}
.client-tabs a{padding:6px 14px;border-radius:7px;font-size:12px;font-weight:700;text-decoration:none;background:#f1f5f9;color:#475569;border:1px solid #e2e8f0}
.client-tabs a.active{background:#4f46e5;color:#fff;border-color:#4f46e5}

.wrap{max-width:1500px;margin:0 auto;padding:20px 24px}
.section-block{padding-top:8px;padding-bottom:4px}
.section-label{display:flex;align-items:center;gap:10px;margin-bottom:14px}
.section-label h2{font-size:17px;font-weight:800;color:#0f172a}
.section-label .pill{font-size:11px;font-weight:700;padding:3px 10px;border-radius:20px;background:#dbeafe;color:#1d4ed8}

/* Stat cards */
.grid-stats{display:grid;grid-template-columns:repeat(6,1fr);gap:12px;margin-bottom:18px}
@media(max-width:1100px){.grid-stats{grid-template-columns:repeat(3,1fr)}}
@media(max-width:700px){.grid-stats{grid-template-columns:repeat(2,1fr)}}
.stat-card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:14px;cursor:pointer;transition:transform .15s,box-shadow .15s;box-shadow:0 1px 6px rgba(0,0,0,.05)}
.stat-card:hover{transform:translateY(-2px);box-shadow:0 6px 20px rgba(0,0,0,.1)}
.stat-icon{width:32px;height:32px;border-radius:10px;display:flex;align-items:center;justify-content:center;margin-bottom:10px;font-size:16px}
.stat-val{font-size:24px;font-weight:900;line-height:1;letter-spacing:-1px}
.stat-lbl{font-size:11px;color:#64748b;font-weight:600;margin-top:4px}
.stat-sub{font-size:10px;font-weight:700;margin-top:3px;color:#94a3b8}
.c-blue .stat-val{color:#2563eb} .c-blue .stat-icon{background:#dbeafe}
.c-green .stat-val{color:#16a34a} .c-green .stat-icon{background:#dcfce7}
.c-purple .stat-val{color:#7c3aed} .c-purple .stat-icon{background:#ede9fe}
.c-orange .stat-val{color:#ea580c} .c-orange .stat-icon{background:#fff7ed}
.c-cyan .stat-val{color:#0891b2} .c-cyan .stat-icon{background:#cffafe}
.c-rose .stat-val{color:#e11d48} .c-rose .stat-icon{background:#fecdd3}
.c-amber .stat-val{color:#d97706} .c-amber .stat-icon{background:#fef3c7}
.c-slate .stat-val{color:#475569} .c-slate .stat-icon{background:#f1f5f9}

/* Cards / tables */
.card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;box-shadow:0 1px 6px rgba(0,0,0,.05);margin-bottom:16px;overflow:hidden}
.card-head{padding:12px 18px;border-bottom:1px solid #f1f5f9;display:flex;align-items:center;justify-content:space-between}
.card-head h3{font-size:13px;font-weight:800;color:#1e293b}
.count-pill{font-size:11px;color:#64748b;background:#f1f5f9;padding:2px 10px;border-radius:20px;font-weight:700}
table{width:100%;border-collapse:collapse;font-size:12px}
th{padding:9px 14px;text-align:left;font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:.7px;border-bottom:1px solid #f1f5f9;background:#fafbfc;white-space:nowrap}
td{padding:10px 14px;border-bottom:1px solid #f8fafc;color:#374151;vertical-align:top}
tr:last-child td{border-bottom:none}
tbody tr:hover td{background:#fafbff}
.clickable{cursor:pointer;color:#2563eb;font-weight:700;text-decoration:underline dotted}
.clickable:hover{color:#1d4ed8}
.no-data{text-align:center;padding:28px;color:#94a3b8;font-size:13px}
.badge{display:inline-flex;padding:2px 8px;border-radius:20px;font-size:10px;font-weight:800}
.b-pos{background:#dcfce7;color:#15803d}
.b-neg{background:#fee2e2;color:#b91c1c}
.b-info{background:#dbeafe;color:#1d4ed8}
.b-warn{background:#fef3c7;color:#92400e}

/* Charts */
.charts-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(360px,1fr));gap:14px;margin-bottom:16px}
.chart-card{background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:14px;box-shadow:0 1px 6px rgba(0,0,0,.05)}
.chart-card h4{font-size:11px;font-weight:700;color:#64748b;margin-bottom:10px}
.chart-wrap{position:relative;height:240px}

/* Modal */
.modal-overlay{position:fixed;inset:0;background:rgba(15,23,42,.6);display:none;align-items:center;justify-content:center;z-index:300;padding:20px}
.modal-overlay.show{display:flex}
.modal{background:#fff;border-radius:14px;max-width:1100px;width:100%;max-height:88vh;overflow:hidden;display:flex;flex-direction:column;box-shadow:0 20px 60px rgba(0,0,0,.4)}
.modal-head{padding:14px 20px;border-bottom:1px solid #e2e8f0;display:flex;align-items:center;justify-content:space-between}
.modal-head h3{font-size:15px;font-weight:800;color:#0f172a}
.modal-close{background:#f1f5f9;border:none;border-radius:8px;width:32px;height:32px;cursor:pointer;font-size:16px;color:#475569}
.modal-close:hover{background:#e2e8f0}
.modal-body{padding:14px 20px;overflow-y:auto}
.lead-row{border-bottom:1px solid #f1f5f9;padding:10px 0}
.lead-row:last-child{border-bottom:none}
.lead-name{font-weight:800;color:#0f172a;font-size:13px}
.lead-meta{font-size:11px;color:#64748b;margin-top:2px}
.lead-msg{background:#f8fafc;border-left:3px solid #6366f1;padding:8px 12px;border-radius:0 6px 6px 0;margin-top:8px;font-size:12px;color:#334155;white-space:pre-wrap}
.lead-msg .mhdr{font-size:10px;color:#94a3b8;font-weight:700;margin-bottom:4px;text-transform:uppercase;letter-spacing:.5px}

.banner{background:#fef9c3;border:1px solid #fde047;border-radius:8px;padding:10px 16px;margin-bottom:14px;font-size:13px;color:#78350f;display:none;align-items:center;gap:10px}
.banner.show{display:flex}
</style>
</head>
<body>

<div class="topbar">
  <div class="logo"><div class="logo-mark">{{client.name[0]}}</div>{{client.name}} — Client Dashboard</div>
  <div style="display:flex;gap:8px;align-items:center">
    <button class="act-btn act-pdf"   id="pdf-btn"   onclick="downloadPDF()">⇩ PDF</button>
    <button class="act-btn act-docx"  id="docx-btn"  onclick="downloadDOCX()">⇩ DOCX</button>
    <button class="act-btn act-share" id="share-btn" onclick="openShareModal()">✉ Share</button>
    <button class="refresh-btn" onclick="loadAll(true)">↻ Refresh</button>
  </div>
</div>

<div class="client-tabs">
  <span class="tlbl">View:</span>
  <a href="/">All Clients</a>
  {% for c in clients %}
  <a href="/client/{{c.slug}}" class="{% if c.slug == client.slug %}active{% endif %}">{{c.name}}</a>
  {% endfor %}
</div>

<div class="wrap">

  <div class="banner" id="banner">⏳ <span id="banner-msg">Loading…</span></div>

  <!-- ─── Date Filter ───────────────────────────────────────────── -->
  <div class="filter-bar" style="background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:12px 18px;margin-bottom:16px;display:flex;align-items:center;gap:10px;flex-wrap:wrap;box-shadow:0 1px 6px rgba(0,0,0,.05)">
    <span style="font-size:10px;font-weight:800;color:#94a3b8;text-transform:uppercase;letter-spacing:1px">Date Range</span>
    <input type="date" id="df-from" style="border:1.5px solid #e2e8f0;border-radius:7px;padding:6px 10px;font-size:12px;background:#f8fafc;font-family:inherit;outline:none">
    <span style="font-size:11px;color:#94a3b8">to</span>
    <input type="date" id="df-to"   style="border:1.5px solid #e2e8f0;border-radius:7px;padding:6px 10px;font-size:12px;background:#f8fafc;font-family:inherit;outline:none">
    <button onclick="applyDate()"     style="padding:6px 14px;background:#4f46e5;color:#fff;border:none;border-radius:7px;font-size:12px;font-weight:700;cursor:pointer">Apply</button>
    <button onclick="presetDate('today')"     style="padding:6px 12px;background:#eef2ff;color:#4338ca;border:1px solid #c7d2fe;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer">Today</button>
    <button onclick="presetDate('yesterday')" style="padding:6px 12px;background:#eef2ff;color:#4338ca;border:1px solid #c7d2fe;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer">Yesterday</button>
    <button onclick="presetDate('last7')"     style="padding:6px 12px;background:#eef2ff;color:#4338ca;border:1px solid #c7d2fe;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer">Last 7 days</button>
    <button onclick="presetDate('last30')"    style="padding:6px 12px;background:#eef2ff;color:#4338ca;border:1px solid #c7d2fe;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer">Last 30 days</button>
    <button onclick="clearDate()"     style="padding:6px 12px;background:#f1f5f9;color:#64748b;border:1px solid #e2e8f0;border-radius:7px;font-size:11px;font-weight:700;cursor:pointer">Clear</button>
    <span id="df-status" style="font-size:11px;color:#6366f1;font-weight:600;margin-left:auto"></span>
  </div>

  <!-- ─── 1. Stats Overview ─────────────────────────────────────── -->
  <div class="section-label"><h2>📊 {{client.name}} — Overview</h2><span class="pill">click any card to drill into leads</span></div>
  <div class="grid-stats" id="stats-grid"></div>

  <!-- ─── 2. Parent Campaigns ───────────────────────────────────── -->
  <div class="section-label"><h2>📧 Parent Campaigns</h2></div>
  <div class="card">
    <div class="card-head"><h3>Campaign Details</h3><span class="count-pill" id="parent-count">0</span></div>
    <div id="parent-table"><div class="no-data">Loading…</div></div>
  </div>

  <!-- ─── 3. Charts ─────────────────────────────────────────────── -->
  <div class="section-label"><h2>📈 Performance Charts</h2></div>
  <div class="charts-grid">
    <div class="chart-card"><h4>Open Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="ch-open"></canvas></div></div>
    <div class="chart-card"><h4>Click Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="ch-click"></canvas></div></div>
    <div class="chart-card"><h4>Reply Rate (%) by Subsequence</h4><div class="chart-wrap"><canvas id="ch-reply"></canvas></div></div>
    <div class="chart-card"><h4>Total Leads by Subsequence</h4><div class="chart-wrap"><canvas id="ch-total"></canvas></div></div>
  </div>

  <!-- ─── 4. Subsequences ───────────────────────────────────────── -->
  <div class="section-label"><h2>🔁 Subsequences</h2></div>
  <div class="card">
    <div class="card-head"><h3>Subsequence Details</h3><span class="count-pill" id="sub-count">0</span></div>
    <div id="sub-table"><div class="no-data">No subsequences for this client.</div></div>
  </div>

  <!-- ─── 5. Positive Leads + Conversations ─────────────────────── -->
  <div class="section-label"><h2>✅ Positive Leads &amp; Conversations</h2></div>
  <div class="card">
    <div class="card-head"><h3>Positive Replies</h3><span class="count-pill" id="pos-count">0</span></div>
    <div id="pos-table"><div class="no-data">No positive leads yet.</div></div>
  </div>

  <!-- ─── 6. LinkedIn Campaigns ─────────────────────────────────── -->
  <div class="section-label"><h2>🔗 LinkedIn — Campaigns</h2><span class="pill">Aimfox</span></div>
  <div class="card">
    <div class="card-head"><h3>Campaigns</h3><span class="count-pill" id="li-camp-count">0</span></div>
    <div id="li-camp-table"><div class="no-data">No LinkedIn campaigns matched.</div></div>
  </div>

  <!-- ─── 7. LinkedIn Conversations ─────────────────────────────── -->
  <div class="card">
    <div class="card-head"><h3>LinkedIn Connection Conversations</h3><span class="count-pill" id="li-conv-count">0</span></div>
    <div id="li-conv-table"><div class="no-data">No conversations yet.</div></div>
  </div>

  <!-- ─── 8. Zoho Sync ──────────────────────────────────────────── -->
  <div class="section-label"><h2>🗂️ Zoho CRM Sync</h2></div>
  <div class="card">
    <div class="card-head">
      <h3>Sync status for {{client.name}}</h3>
      <button class="refresh-btn" style="background:#4f46e5;color:#fff;border:none" onclick="syncNow()">🔄 Sync Now</button>
    </div>
    <div id="zoho-body" style="padding:14px 18px;color:#475569;font-size:13px">Loading…</div>
  </div>

</div>

<!-- Lead modal -->
<div class="modal-overlay" id="modal-ovl" onclick="if(event.target===this)closeModal()">
  <div class="modal">
    <div class="modal-head">
      <h3 id="modal-title">Leads</h3>
      <button class="modal-close" onclick="closeModal()">✕</button>
    </div>
    <div class="modal-body" id="modal-body"></div>
  </div>
</div>

<!-- Share email modal -->
<div class="modal-overlay" id="share-ovl" onclick="if(event.target===this)closeShareModal()">
  <div class="modal" style="max-width:480px">
    <div class="modal-head">
      <h3>✉ Share {{client.name}} Report via Email</h3>
      <button class="modal-close" onclick="closeShareModal()">✕</button>
    </div>
    <div class="modal-body">
      <p style="font-size:12px;color:#64748b;margin-bottom:12px">Send the {{client.name}} PDF analytics report to:</p>
      <input type="email" id="share-email-inp" placeholder="someone@example.com"
             style="width:100%;padding:10px 12px;border:1.5px solid #e2e8f0;border-radius:8px;font-size:13px;font-family:inherit;outline:none;margin-bottom:12px">
      <div id="share-msg" style="display:none;font-size:12px;font-weight:600;padding:8px 10px;border-radius:6px;margin-bottom:10px"></div>
      <div style="display:flex;gap:8px;justify-content:flex-end">
        <button class="act-btn" style="background:#f1f5f9;color:#475569" onclick="closeShareModal()">Cancel</button>
        <button class="act-btn act-share" id="share-send-btn" onclick="sendShareEmail()">Send Report</button>
      </div>
    </div>
  </div>
</div>

<script>
const SLUG = "{{client.slug}}";
const CLIENT = "{{client.name}}";
let DATA = null;       // currently displayed (may be filtered)
let DATA_FULL = null;  // unfiltered original from /api/client/<slug>
let CHARTS = {};

const POSITIVE = new Set(["interested","meeting booked","positive","meeting request","will buy","warm","demo request"]);

function fmt(n){return (typeof n==="number"?n:0).toLocaleString()}
function pct(num,den){if(!den) return "0%"; return (num/den*100).toFixed(1)+"%"}
function rateColor(p){return p>=20?"#16a34a":p>=10?"#ca8a04":"#dc2626"}

async function loadAll(force){
  showBanner("Loading…");
  try{
    const r = await fetch(`/api/client/${SLUG}${force?"?_=".concat(Date.now()):""}`);
    const d = await r.json();
    if(d.error) throw new Error(d.error);
    DATA_FULL = d;
    DATA = d;
    if(d.loading){
      showBanner("⏳ Smartlead data still fetching in background — showing partial data…");
      setTimeout(()=>loadAll(true), 12000);
    } else {
      hideBanner();
    }
    // Re-apply current date filter (if any) to fresh data
    if(document.getElementById("df-from").value || document.getElementById("df-to").value){
      applyDate(true);
    } else {
      renderAll();
    }
  } catch(e){
    showBanner("Error: " + e.message);
  }
}

// ── Date filter helpers ──────────────────────────────────────────
function _todayStr(d){ d = d || new Date(); return d.toISOString().slice(0,10); }
function presetDate(kind){
  const today = new Date();
  let from, to;
  if(kind==="today"){ from = to = _todayStr(today); }
  else if(kind==="yesterday"){
    const y = new Date(today); y.setDate(y.getDate()-1);
    from = to = _todayStr(y);
  }
  else if(kind==="last7"){
    const a = new Date(today); a.setDate(a.getDate()-6);
    from = _todayStr(a); to = _todayStr(today);
  }
  else if(kind==="last30"){
    const a = new Date(today); a.setDate(a.getDate()-29);
    from = _todayStr(a); to = _todayStr(today);
  }
  document.getElementById("df-from").value = from;
  document.getElementById("df-to").value   = to;
  applyDate();
}
function clearDate(){
  document.getElementById("df-from").value = "";
  document.getElementById("df-to").value   = "";
  document.getElementById("df-status").textContent = "";
  DATA = DATA_FULL;
  renderAll();
}

function applyDate(silent){
  const from = document.getElementById("df-from").value;
  const to   = document.getElementById("df-to").value;
  if(!from && !to){ clearDate(); return; }
  if(!DATA_FULL){ return; }
  const lo = from ? from + "T00:00:00" : "0000-01-01";
  const hi = to   ? to   + "T23:59:59" : "9999-12-31";

  const inRange = ts => {
    if(!ts) return false;
    const t = String(ts).slice(0,19);
    return t >= lo && t <= hi;
  };
  // A lead is "in range" if its sent_time falls in the window
  // (campaigns where no leads were sent in window get total=0)
  const filterCamp = (camp) => {
    const leads = (camp.leads||[]).filter(l => inRange(l.sent_time));
    const total = leads.length;
    const opened       = leads.filter(l => inRange(l.open_time)).length;
    const clicked      = leads.filter(l => inRange(l.click_time)).length;
    const replied      = leads.filter(l => inRange(l.reply_time)).length;
    const bounced      = leads.filter(l => (l.category||"").toLowerCase()==="bounced").length;
    const unsubscribed = leads.filter(l => (l.category||"").toLowerCase()==="unsubscribed").length;
    const added_to_sub = camp.added_to_sub_emails
        ? leads.filter(l => (camp.added_to_sub_emails||[]).includes(l.email)).length
        : leads.filter(l => l.added_to_sub).length;
    return {
      ...camp,
      leads, total, opened, clicked, replied, bounced, unsubscribed, added_to_sub,
      open_rate:  total ? +(opened/total*100).toFixed(1)  : 0,
      reply_rate: total ? +(replied/total*100).toFixed(1) : 0,
      click_rate: total ? +(clicked/total*100).toFixed(1) : 0,
    };
  };

  const filtered = JSON.parse(JSON.stringify(DATA_FULL));
  filtered.parent_campaigns = (DATA_FULL.parent_campaigns||[])
      .map(filterCamp)
      .filter(c => (c.total||0) > 0 || (c.leads||[]).length > 0);  // hide pure-zero rows
  filtered.subsequences = (DATA_FULL.subsequences||[])
      .map(filterCamp)
      .filter(s => (s.total||0) > 0);

  // Positive leads — keep those whose reply_time falls in window
  const POS = new Set(["interested","meeting booked","positive","meeting request","will buy","warm","demo request"]);
  filtered.positive_leads = [];
  filtered.parent_campaigns.forEach(p=>{
    (p.leads||[]).forEach(l=>{
      if(POS.has((l.category||"").toLowerCase())){
        filtered.positive_leads.push({
          ...l,
          name: l.name || l.email || "Unknown",
          campaign: p.raw_name || p.client || "",
          messages: l.messages || [],
        });
      }
    });
  });

  // Recompute summary from filtered totals
  const sum = (arr,k) => arr.reduce((a,r)=>a+(r[k]||0), 0);
  filtered.stats = {
    sent:         sum(filtered.parent_campaigns, "total"),
    opened:       sum(filtered.parent_campaigns, "opened"),
    clicked:      sum(filtered.parent_campaigns, "clicked"),
    replied:      sum(filtered.parent_campaigns, "replied"),
    bounced:      sum(filtered.parent_campaigns, "bounced"),
    unsubscribed: sum(filtered.parent_campaigns, "unsubscribed"),
    added_to_sub: sum(filtered.parent_campaigns, "added_to_sub"),
    sub_total:    sum(filtered.subsequences,    "total"),
    sub_opened:   sum(filtered.subsequences,    "opened"),
    sub_clicked:  sum(filtered.subsequences,    "clicked"),
    sub_replied:  sum(filtered.subsequences,    "replied"),
    positive:     filtered.positive_leads.length,
  };

  DATA = filtered;
  const lbl = (from && to && from === to) ? from
            : (from && to)               ? `${from} to ${to}`
            : (from)                     ? `from ${from}`
            : `until ${to}`;
  document.getElementById("df-status").textContent =
      `Filtered ${lbl} • ${filtered.parent_campaigns.length} campaigns • ${filtered.stats.sent.toLocaleString()} sent`;
  renderAll();
}

function showBanner(msg){
  const b = document.getElementById("banner");
  document.getElementById("banner-msg").textContent = msg;
  b.classList.add("show");
}
function hideBanner(){document.getElementById("banner").classList.remove("show")}

function renderAll(){
  renderStats();
  renderParents();
  renderCharts();
  renderSubs();
  renderPositive();
  renderLI();
  renderZoho();
}

function renderStats(){
  const s = DATA.stats || {};
  const cards = [
    {key:"sent",         lbl:"Sent",          val:s.sent,         icon:"📤", col:"c-blue",
     leads:()=>collectLeads(p=>true)},
    {key:"opened",       lbl:"Opened",        val:s.opened,       icon:"👁",  col:"c-green",
     sub:pct(s.opened,s.sent),
     leads:()=>collectLeads(l=>l.open_time)},
    {key:"clicked",      lbl:"Clicked",       val:s.clicked,      icon:"🖱",  col:"c-cyan",
     sub:pct(s.clicked,s.sent),
     leads:()=>collectLeads(l=>l.click_time)},
    {key:"replied",      lbl:"Replied",       val:s.replied,      icon:"💬", col:"c-amber",
     sub:pct(s.replied,s.sent),
     leads:()=>collectLeads(l=>l.reply_time)},
    {key:"added_to_sub", lbl:"Added to Sub",  val:s.added_to_sub, icon:"➕", col:"c-purple",
     sub:pct(s.added_to_sub,s.sent),
     leads:()=>subLeads(()=>true)},
    {key:"sub_opened",   lbl:"Sub Opened",    val:s.sub_opened,   icon:"📖", col:"c-orange",
     sub:pct(s.sub_opened,s.sub_total),
     leads:()=>subLeads(l=>l.open_time)},
    {key:"sub_clicked",  lbl:"Sub Clicked",   val:s.sub_clicked,  icon:"🎯", col:"c-rose",
     sub:pct(s.sub_clicked,s.sub_total),
     leads:()=>subLeads(l=>l.click_time)},
    {key:"sub_replied",  lbl:"Sub Replied",   val:s.sub_replied,  icon:"📨", col:"c-amber",
     sub:pct(s.sub_replied,s.sub_total),
     leads:()=>subLeads(l=>l.reply_time)},
    {key:"positive",     lbl:"Positive Replies", val:s.positive,  icon:"✅", col:"c-green",
     leads:()=>(DATA.positive_leads||[])},
    {key:"bounced",      lbl:"Bounced",       val:s.bounced,      icon:"⚠",  col:"c-rose",
     leads:()=>collectLeads(l=>(l.category||"").toLowerCase()==="bounced")},
    {key:"unsubscribed", lbl:"Unsubscribed",  val:s.unsubscribed, icon:"🚫", col:"c-slate",
     leads:()=>collectLeads(l=>(l.category||"").toLowerCase()==="unsubscribed")},
  ];
  document.getElementById("stats-grid").innerHTML = cards.map((c,i)=>`
    <div class="stat-card ${c.col}" data-idx="${i}">
      <div class="stat-icon">${c.icon}</div>
      <div class="stat-val">${fmt(c.val)}</div>
      <div class="stat-lbl">${c.lbl}</div>
      ${c.sub?`<div class="stat-sub">${c.sub}</div>`:""}
    </div>`).join("");
  // Wire click handlers
  document.querySelectorAll("#stats-grid .stat-card").forEach((el,i)=>{
    el.onclick = ()=>openLeadModal(`${cards[i].lbl} — ${CLIENT}`, cards[i].leads());
  });
}

function collectLeads(filter){
  const out = [];
  (DATA.parent_campaigns||[]).forEach(p=>{
    (p.leads||[]).forEach(l=>{ if(filter(l)) out.push({...l, _campaign:p.raw_name||p.client}); });
  });
  return out;
}
function subLeads(filter){
  const out = [];
  (DATA.subsequences||[]).forEach(s=>{
    (s.leads||[]).forEach(l=>{ if(filter(l)) out.push({...l, _campaign:`${s.parent} → ${s.subsequence}`}); });
  });
  return out;
}

function renderParents(){
  const rows = DATA.parent_campaigns||[];
  document.getElementById("parent-count").textContent = rows.length;
  if(!rows.length){document.getElementById("parent-table").innerHTML='<div class="no-data">No parent campaigns for this client.</div>';return;}
  const html = `<table><thead><tr>
    <th>Campaign</th><th>Status</th>
    <th style="text-align:right">Total</th><th style="text-align:right">Opened</th>
    <th style="text-align:right">Clicked</th><th style="text-align:right">Replied</th>
    <th style="text-align:right">Bounced</th><th style="text-align:right">Added to Sub</th>
    <th style="text-align:right">Open %</th><th style="text-align:right">Reply %</th>
  </tr></thead><tbody>
  ${rows.map((r,i)=>`<tr>
    <td>${escapeHtml(r.raw_name || r.client || '')}</td>
    <td><span class="badge b-info">${escapeHtml(r.status||'—')}</span></td>
    <td style="text-align:right" class="clickable" onclick="openParentLeads(${i},'all','All leads')">${fmt(r.total)}</td>
    <td style="text-align:right" class="clickable" onclick="openParentLeads(${i},'open','Opened')">${fmt(r.opened)}</td>
    <td style="text-align:right" class="clickable" onclick="openParentLeads(${i},'click','Clicked')">${fmt(r.clicked)}</td>
    <td style="text-align:right" class="clickable" onclick="openParentLeads(${i},'reply','Replied')">${fmt(r.replied)}</td>
    <td style="text-align:right">${fmt(r.bounced)}</td>
    <td style="text-align:right" class="clickable" onclick="openParentLeads(${i},'all','Added to Sub')">${fmt(r.added_to_sub)}</td>
    <td style="text-align:right;color:${rateColor(r.open_rate||0)};font-weight:700">${(+r.open_rate||0).toFixed(1)}%</td>
    <td style="text-align:right;color:${rateColor(r.reply_rate||0)};font-weight:700">${(+r.reply_rate||0).toFixed(1)}%</td>
  </tr>`).join("")}
  </tbody></table>`;
  document.getElementById("parent-table").innerHTML = html;
}

function openParentLeads(idx, kind, title){
  const p = (DATA.parent_campaigns||[])[idx];
  if(!p) return;
  let leads = p.leads || [];
  if(kind==="open")  leads = leads.filter(l=>l.open_time);
  if(kind==="click") leads = leads.filter(l=>l.click_time);
  if(kind==="reply") leads = leads.filter(l=>l.reply_time);
  openLeadModal(`${title} — ${escapeHtml(p.raw_name||p.client)}`, leads);
}

function renderCharts(){
  const subs = (DATA.subsequences||[]).slice().sort((a,b)=>(b.total||0)-(a.total||0)).slice(0,12);
  const labels = subs.map(s=>truncate(s.subsequence||"",18));
  drawBar("ch-open",  labels, subs.map(s=>+(s.open_rate||0).toFixed(1)),  "#16a34a", "Open %");
  drawBar("ch-click", labels, subs.map(s=>+(s.click_rate||0).toFixed(1)), "#0891b2", "Click %");
  drawBar("ch-reply", labels, subs.map(s=>+(s.reply_rate||0).toFixed(1)), "#d97706", "Reply %");
  drawBar("ch-total", labels, subs.map(s=>+(s.total||0)),                 "#7c3aed", "Total Leads");
}
function drawBar(id, labels, data, color, lbl){
  const ctx = document.getElementById(id);
  if(!ctx) return;
  if(CHARTS[id]) CHARTS[id].destroy();
  CHARTS[id] = new Chart(ctx, {
    type:'bar',
    data:{labels, datasets:[{label:lbl, data, backgroundColor:color+"cc", borderColor:color, borderWidth:1, borderRadius:4}]},
    options:{maintainAspectRatio:false, plugins:{legend:{display:false}}, scales:{x:{ticks:{font:{size:9}}}, y:{beginAtZero:true, ticks:{font:{size:10}}}}}
  });
}

function renderSubs(){
  const rows = DATA.subsequences||[];
  document.getElementById("sub-count").textContent = rows.length;
  if(!rows.length) return;
  document.getElementById("sub-table").innerHTML = `<table><thead><tr>
    <th>Subsequence</th>
    <th style="text-align:right">Total</th><th style="text-align:right">Opened</th>
    <th style="text-align:right">Clicked</th><th style="text-align:right">Replied</th>
    <th style="text-align:right">Open %</th><th style="text-align:right">Reply %</th>
  </tr></thead><tbody>
  ${rows.map((r,i)=>`<tr>
    <td>${escapeHtml(r.subsequence||'')}</td>
    <td style="text-align:right" class="clickable" onclick="openSubLeads(${i},'all','All')">${fmt(r.total)}</td>
    <td style="text-align:right" class="clickable" onclick="openSubLeads(${i},'open','Opened')">${fmt(r.opened)}</td>
    <td style="text-align:right" class="clickable" onclick="openSubLeads(${i},'click','Clicked')">${fmt(r.clicked)}</td>
    <td style="text-align:right" class="clickable" onclick="openSubLeads(${i},'reply','Replied')">${fmt(r.replied)}</td>
    <td style="text-align:right;color:${rateColor(r.open_rate||0)};font-weight:700">${(+r.open_rate||0).toFixed(1)}%</td>
    <td style="text-align:right;color:${rateColor(r.reply_rate||0)};font-weight:700">${(+r.reply_rate||0).toFixed(1)}%</td>
  </tr>`).join("")}
  </tbody></table>`;
}
function openSubLeads(idx, kind, title){
  const s = (DATA.subsequences||[])[idx]; if(!s) return;
  let leads = s.leads || [];
  if(kind==="open")  leads = leads.filter(l=>l.open_time);
  if(kind==="click") leads = leads.filter(l=>l.click_time);
  if(kind==="reply") leads = leads.filter(l=>l.reply_time);
  openLeadModal(`${title} — ${escapeHtml(s.subsequence||'')}`, leads);
}

function renderPositive(){
  const rows = DATA.positive_leads||[];
  document.getElementById("pos-count").textContent = rows.length;
  if(!rows.length){return;}
  document.getElementById("pos-table").innerHTML = `<table><thead><tr>
    <th>Lead</th><th>Email</th><th>Campaign</th><th>Category</th><th>Reply Date</th><th>Conversation</th>
  </tr></thead><tbody>
  ${rows.map((l,i)=>`<tr>
    <td><strong>${escapeHtml(l.name||'')}</strong></td>
    <td style="font-size:11px;color:#64748b">${escapeHtml(l.email||'')}</td>
    <td style="font-size:11px">${escapeHtml(l.campaign||'')}</td>
    <td><span class="badge b-pos">${escapeHtml(l.category||'')}</span></td>
    <td style="font-size:11px">${escapeHtml((l.reply_time||'').slice(0,10))}</td>
    <td><span class="clickable" onclick="openConvo(${i})">View ${(l.messages||[]).length} msgs →</span></td>
  </tr>`).join("")}
  </tbody></table>`;
}
function openConvo(i){
  const l = (DATA.positive_leads||[])[i]; if(!l) return;
  const msgs = (l.messages||[]);
  const body = msgs.length
    ? msgs.map(m=>`<div class="lead-msg"><div class="mhdr">${escapeHtml(m.type||"")} · ${escapeHtml(m.from||"")} · ${escapeHtml(m.time||"")}</div>${escapeHtml(m.body||"")}</div>`).join("")
    : `<div class="no-data">No conversation messages stored.</div>`;
  document.getElementById("modal-title").textContent = `Conversation — ${l.name}`;
  document.getElementById("modal-body").innerHTML = `
    <div class="lead-row">
      <div class="lead-name">${escapeHtml(l.name||"")}</div>
      <div class="lead-meta">${escapeHtml(l.email||"")} · ${escapeHtml(l.campaign||"")} · <span class="badge b-pos">${escapeHtml(l.category||"")}</span></div>
    </div>
    ${body}`;
  document.getElementById("modal-ovl").classList.add("show");
}

function renderLI(){
  const li = DATA.linkedin || {};
  const camps = li.campaigns || [];
  document.getElementById("li-camp-count").textContent = camps.length;
  if(camps.length){
    document.getElementById("li-camp-table").innerHTML = `<table><thead><tr>
      <th>Campaign</th><th>Type</th><th>State</th><th>Owner</th>
      <th style="text-align:right">Targets</th><th style="text-align:right">Accepted</th><th style="text-align:right">Replies</th><th style="text-align:right">Done %</th>
    </tr></thead><tbody>
    ${camps.map(c=>`<tr>
      <td>${escapeHtml(c.name||'')}</td>
      <td><span class="badge b-info">${escapeHtml(c.type||'—')}</span></td>
      <td><span class="badge ${c.state==='ACTIVE'?'b-pos':'b-warn'}">${escapeHtml(c.state||'')}</span></td>
      <td style="font-size:11px">${escapeHtml((c.owners||[]).join(", "))}</td>
      <td style="text-align:right">${fmt(c.targets)}</td>
      <td style="text-align:right;color:#16a34a;font-weight:700">${fmt(c.accepted)}</td>
      <td style="text-align:right;color:#d97706;font-weight:700">${fmt(c.replies)}</td>
      <td style="text-align:right">${(+c.completion_pct||0).toFixed(0)}%</td>
    </tr>`).join("")}
    </tbody></table>`;
  }
  const convs = li.conversations || [];
  document.getElementById("li-conv-count").textContent = convs.length;
  if(convs.length){
    document.getElementById("li-conv-table").innerHTML = `<table><thead><tr>
      <th>Lead</th><th>Occupation</th><th>Campaign</th><th>Account</th>
      <th style="text-align:right">Msgs</th><th style="text-align:right">Unread</th><th>Last Message</th>
    </tr></thead><tbody>
    ${convs.map(c=>`<tr>
      <td><strong>${escapeHtml(c.lead||'')}</strong></td>
      <td style="font-size:11px;color:#64748b">${escapeHtml(c.occupation||'')}</td>
      <td style="font-size:11px">${escapeHtml(c.campaign||'')}</td>
      <td style="font-size:11px">${escapeHtml(c.account||'')}</td>
      <td style="text-align:right">${fmt(c.msgs)}</td>
      <td style="text-align:right;color:${c.unread>0?'#dc2626':'#94a3b8'};font-weight:700">${fmt(c.unread)}</td>
      <td style="font-size:11px">${escapeHtml((c.last||'').slice(0,16))}</td>
    </tr>`).join("")}
    </tbody></table>`;
  }
}

function renderZoho(){
  const z = DATA.zoho || {};
  document.getElementById("zoho-body").innerHTML = `
    <div style="display:flex;gap:24px;flex-wrap:wrap">
      <div><div style="font-size:11px;color:#94a3b8;font-weight:700">LAST SYNC</div><div style="font-size:14px;font-weight:700">${escapeHtml(z.last_sync || '— never —')}</div></div>
      <div><div style="font-size:11px;color:#94a3b8;font-weight:700">TOTAL SYNCED</div><div style="font-size:18px;font-weight:900;color:#1e293b">${fmt(z.total)}</div></div>
      <div><div style="font-size:11px;color:#94a3b8;font-weight:700">CREATED</div><div style="font-size:18px;font-weight:900;color:#16a34a">${fmt(z.created)}</div></div>
      <div><div style="font-size:11px;color:#94a3b8;font-weight:700">UPDATED</div><div style="font-size:18px;font-weight:900;color:#2563eb">${fmt(z.updated)}</div></div>
    </div>`;
}

async function syncNow(){
  const body = document.getElementById("zoho-body");
  body.innerHTML = '<div style="color:#6366f1">⏳ Syncing to Zoho…</div>';
  try{
    const r = await fetch("/api/zoho-sync", {method:"POST"});
    const d = await r.json();
    if(d.ok){
      body.innerHTML = `<div style="color:#16a34a;font-weight:700">✅ Sync complete: ${fmt(d.created||0)} created, ${fmt(d.updated||0)} updated</div>`;
      setTimeout(()=>loadAll(true), 1500);
    } else {
      body.innerHTML = `<div style="color:#dc2626">❌ ${escapeHtml(d.error||'Sync failed')}</div>`;
    }
  } catch(e){
    body.innerHTML = `<div style="color:#dc2626">❌ ${escapeHtml(e.message)}</div>`;
  }
}

function openLeadModal(title, leads){
  document.getElementById("modal-title").textContent = `${title} (${leads.length})`;
  if(!leads.length){
    document.getElementById("modal-body").innerHTML = '<div class="no-data">No leads to show.</div>';
  } else {
    document.getElementById("modal-body").innerHTML = leads.slice(0, 200).map(l=>{
      const msgs = (l.messages||[]).slice(0,3).map(m=>
        `<div class="lead-msg"><div class="mhdr">${escapeHtml(m.type||"")} · ${escapeHtml(m.time||"")}</div>${escapeHtml((m.body||"").slice(0,300))}</div>`
      ).join("");
      const cat = l.category ? `<span class="badge ${POSITIVE.has((l.category||"").toLowerCase())?'b-pos':'b-info'}">${escapeHtml(l.category)}</span>` : "";
      return `<div class="lead-row">
        <div class="lead-name">${escapeHtml(l.name||l.email||"Unknown")} ${cat}</div>
        <div class="lead-meta">${escapeHtml(l.email||"")} · ${escapeHtml(l._campaign||l.campaign||"")} · ${l.reply_time?`replied ${escapeHtml(l.reply_time.slice(0,10))}`:l.open_time?`opened ${escapeHtml(l.open_time.slice(0,10))}`:""}</div>
        ${msgs}
      </div>`;
    }).join("") + (leads.length>200?`<div class="no-data">Showing first 200 of ${leads.length}</div>`:"");
  }
  document.getElementById("modal-ovl").classList.add("show");
}
function closeModal(){document.getElementById("modal-ovl").classList.remove("show")}
function escapeHtml(s){return String(s||"").replace(/[&<>"']/g, c=>({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]))}
function truncate(s,n){return (s||"").length>n? (s||"").slice(0,n-1)+"…":(s||"")}

document.addEventListener("keydown", e=>{ if(e.key==="Escape"){ closeModal(); closeShareModal(); }});

// ── Download / Share (per-client filtered) ─────────────────────
function _getFilteredPayload(){
  return {
    client:           CLIENT,
    parent_analytics: DATA?.parent_campaigns || [],
    sub_analytics:    DATA?.subsequences     || [],
    summary:          DATA?.stats            || {},
    positive_leads:   DATA?.positive_leads   || [],
  };
}

async function downloadPDF(){
  const btn = document.getElementById("pdf-btn");
  btn.disabled = true; btn.textContent = "⏳ PDF…";
  try{
    const r = await fetch("/api/export/pdf", {
      method:"POST", headers:{"Content-Type":"application/json"},
      body: JSON.stringify(_getFilteredPayload()),
    });
    if(!r.ok){ const d = await r.json().catch(()=>({})); throw new Error(d.error || "HTTP "+r.status); }
    const blob = await r.blob();
    const url  = URL.createObjectURL(blob);
    const a = document.createElement("a");
    a.href = url; a.download = `${SLUG}_analytics_${new Date().toISOString().slice(0,10)}.pdf`;
    a.click(); URL.revokeObjectURL(url);
  } catch(e){
    alert("PDF failed: " + e.message);
  } finally {
    btn.disabled = false; btn.innerHTML = "⇩ PDF";
  }
}

async function downloadDOCX(){
  const btn = document.getElementById("docx-btn");
  btn.disabled = true; btn.textContent = "⏳ DOCX…";
  try{
    const r = await fetch("/api/export/docx", {
      method:"POST", headers:{"Content-Type":"application/json"},
      body: JSON.stringify(_getFilteredPayload()),
    });
    if(!r.ok){ const d = await r.json().catch(()=>({})); throw new Error(d.error || "HTTP "+r.status); }
    const blob = await r.blob();
    const url  = URL.createObjectURL(blob);
    const a = document.createElement("a");
    a.href = url; a.download = `${SLUG}_analytics_${new Date().toISOString().slice(0,10)}.docx`;
    a.click(); URL.revokeObjectURL(url);
  } catch(e){
    alert("DOCX failed: " + e.message);
  } finally {
    btn.disabled = false; btn.innerHTML = "⇩ DOCX";
  }
}

function openShareModal(){
  document.getElementById("share-ovl").classList.add("show");
  document.getElementById("share-email-inp").focus();
}
function closeShareModal(){
  document.getElementById("share-ovl").classList.remove("show");
  const m = document.getElementById("share-msg"); m.style.display="none"; m.textContent="";
  document.getElementById("share-email-inp").value="";
}

async function sendShareEmail(){
  const email = document.getElementById("share-email-inp").value.trim();
  const msgEl = document.getElementById("share-msg");
  const sendBtn = document.getElementById("share-send-btn");
  if(!email){
    msgEl.textContent = "Please enter an email address.";
    msgEl.style.cssText = "display:block;background:#fee2e2;color:#b91c1c;padding:8px 10px;border-radius:6px;margin-bottom:10px;font-size:12px;font-weight:600";
    return;
  }
  sendBtn.disabled = true; sendBtn.textContent = "Sending…";
  msgEl.textContent = "Generating PDF and sending…";
  msgEl.style.cssText = "display:block;background:#eef2ff;color:#3730a3;padding:8px 10px;border-radius:6px;margin-bottom:10px;font-size:12px;font-weight:600";
  try{
    const r = await fetch("/api/share/email", {
      method:"POST", headers:{"Content-Type":"application/json"},
      body: JSON.stringify({email, ..._getFilteredPayload()}),
    });
    const d = await r.json();
    if(!d.ok) throw new Error(d.error || "Unknown error");
    msgEl.textContent = `✓ Report sent to ${email}`;
    msgEl.style.cssText = "display:block;background:#dcfce7;color:#15803d;padding:8px 10px;border-radius:6px;margin-bottom:10px;font-size:12px;font-weight:600";
  } catch(e){
    msgEl.textContent = "Failed: " + e.message;
    msgEl.style.cssText = "display:block;background:#fee2e2;color:#b91c1c;padding:8px 10px;border-radius:6px;margin-bottom:10px;font-size:12px;font-weight:600";
  } finally {
    sendBtn.disabled = false; sendBtn.textContent = "Send Report";
  }
}

loadAll();
</script>
</body>
</html>"""


# ── Background scheduler: keeps cache fresh so dashboard is always instant ──
def _scheduled_refresher():
    """Refreshes email cache every REFRESH_EVERY seconds (25 min by default).
    Runs forever in a daemon thread. First refresh waits REFRESH_EVERY since
    startup pre-fetch already covers the initial load."""
    while True:
        try:
            time.sleep(REFRESH_EVERY)
            log.info("[scheduler] Auto-refreshing email cache...")
            t0 = time.time()
            get_email_data(force=True)
            log.info("[scheduler] Auto-refresh done in %.1fs", time.time() - t0)
        except Exception as e:
            log.error("[scheduler] Refresh failed: %s", e, exc_info=True)
            time.sleep(60)  # back off briefly on error


# Pre-populate email cache in background so it's ready when the page first loads.
# Runs under both `python unified_dashboard.py` and gunicorn (Render).
if os.getenv("SMARTLEAD_API_KEY") and not os.getenv("DISABLE_SCHEDULER"):
    threading.Thread(target=get_email_data, daemon=True).start()
    threading.Thread(target=_scheduled_refresher, daemon=True).start()
    log.info("[scheduler] Auto-refresh enabled (every %d min)", REFRESH_EVERY // 60)


if __name__ == "__main__":
    log.info("Unified Analytics Dashboard → http://localhost:%d", PORT)
    app.run(host="0.0.0.0", port=PORT, debug=False)
